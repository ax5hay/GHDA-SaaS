#!/usr/bin/env python3
"""
Government Health Data Automation - Post-POC MVP
Single-file analyzer that produces award-winning granular analysis

This script uses Claude AI to handle completely unpredictable ground worker input
and produces comprehensive, government-contract-worthy analysis reports.

Usage:
    python poc_analyzer.py path/to/report.docx

Requirements:
    pip install anthropic python-docx PyPDF2 python-dateutil
    Set environment variable: ANTHROPIC_API_KEY
"""

import os
import sys
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

# Document processing imports
try:
    import docx
    from PyPDF2 import PdfReader
except ImportError:
    print("ERROR: Required packages not installed.")
    print("Install with: pip install anthropic python-docx PyPDF2 python-dateutil")
    sys.exit(1)

# AI import
try:
    import anthropic
except ImportError:
    print("ERROR: Anthropic package not installed.")
    print("Install with: pip install anthropic")
    sys.exit(1)


class GovernmentHealthAnalyzer:
    """
    Post-POC MVP: AI-powered health report analyzer

    Handles completely unpredictable input from ground workers and produces
    government-contract-worthy granular analysis.
    """

    def __init__(self, api_key: Optional[str] = None):
        """Initialize analyzer with Claude AI."""
        self.api_key = api_key or os.environ.get("ANTHROPIC_API_KEY")
        if not self.api_key:
            raise ValueError(
                "ANTHROPIC_API_KEY environment variable must be set. "
                "Get your API key from https://console.anthropic.com/"
            )

        self.client = anthropic.Anthropic(api_key=self.api_key)
        self.model = "claude-sonnet-4-20250514"  # Latest model

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                text_parts.append(" | ".join(row_data))

        return "\n".join(text_parts)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)

        return "\n".join(text_parts)

    def extract_text(self, file_path: str) -> str:
        """Extract text from document (DOCX or PDF)."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension == ".docx":
            return self.extract_text_from_docx(str(file_path))
        elif extension == ".pdf":
            return self.extract_text_from_pdf(str(file_path))
        else:
            # Try reading as plain text
            return file_path.read_text(encoding="utf-8", errors="ignore")

    def analyze_document(self, raw_text: str) -> Dict[str, Any]:
        """
        AI-powered document analysis.

        Extracts structured data and performs deep analysis on completely
        unpredictable ground worker input.
        """

        analysis_prompt = f"""You are an expert government health data analyst specializing in Preconception and Maternal Health Clinics (PPC) in India.

You will analyze a field report written by a ground health worker. The input may be:
- In broken English, Hinglish, or Roman Hindi
- Poorly formatted with spelling errors
- Incomplete or vague
- Using local slang or colloquialisms

Your task is to extract ALL possible information and perform DEEP, GRANULAR analysis that would impress government stakeholders and justify a major contract.

DOCUMENT TEXT:
{raw_text}

---

Provide a comprehensive JSON analysis with the following structure:

{{
  "document_metadata": {{
    "extraction_timestamp": "ISO timestamp",
    "document_quality_assessment": {{
      "readability_score": 0-100,
      "completeness_score": 0-100,
      "data_clarity": "poor/fair/good/excellent",
      "language_detected": "english/hinglish/hindi/mixed",
      "special_challenges": ["list of data quality issues"]
    }}
  }},

  "facility_information": {{
    "facility_name": "extracted name or 'Unknown'",
    "facility_type": "CHC/PHC/Sub-Center/District Hospital/Unknown",
    "location": {{
      "block": "extracted or 'Unknown'",
      "district": "extracted or 'Unknown'",
      "state": "extracted or 'Unknown'"
    }},
    "confidence_scores": {{
      "facility_name": 0.0-1.0,
      "location": 0.0-1.0
    }},
    "extraction_notes": "how this information was inferred"
  }},

  "clinic_details": {{
    "clinic_date": "YYYY-MM-DD or null",
    "day_of_week": "Monday-Sunday or null",
    "timing": {{
      "start_time": "HH:MM or null",
      "end_time": "HH:MM or null",
      "duration_hours": float or null
    }},
    "clinic_type": "Regular/Special/Outreach/Camp/Unknown",
    "reporting_staff": {{
      "name": "extracted or 'Unknown'",
      "designation": "extracted or 'Unknown'",
      "contact": "extracted or null"
    }},
    "confidence_scores": {{
      "date": 0.0-1.0,
      "timing": 0.0-1.0
    }}
  }},

  "infrastructure_assessment": {{
    "dedicated_space": {{
      "available": true/false/null,
      "type": "single_room/multiple_rooms/shared_space/none/unknown",
      "privacy_ensured": true/false/null,
      "issues_identified": ["list of infrastructure problems mentioned"]
    }},
    "waiting_area": {{
      "available": true/false/null,
      "adequacy": "inadequate/barely_adequate/adequate/good/unknown"
    }},
    "signage_and_awareness": {{
      "ppc_board_displayed": true/false/null,
      "beneficiary_awareness": "low/medium/high/unknown"
    }},
    "critical_gaps": [
      {{
        "gap_type": "space/equipment/supplies/staffing/cleanliness/other",
        "severity": "low/medium/high/critical",
        "description": "detailed description",
        "evidence_from_text": "exact quote or paraphrase",
        "potential_impact": "analysis of how this affects service delivery"
      }}
    ],
    "overall_infrastructure_score": 0-100
  }},

  "beneficiary_attendance_analysis": {{
    "expected_beneficiaries": int or null,
    "actual_attendance": int or null,
    "attendance_rate": float or null,
    "attendance_quality": "critical_failure/poor/fair/good/excellent/unknown",

    "individual_beneficiaries": [
      {{
        "beneficiary_identifier": "name/id if mentioned",
        "attended": true/false,
        "age": int or null,
        "gestational_status": "preconception/pregnant/postnatal/unknown",
        "reason_for_absence": "detailed reason if absent",
        "family_situation": "any family context mentioned"
      }}
    ],

    "attendance_barriers_detected": [
      {{
        "barrier_type": "communication/cultural/economic/health/logistical/systemic/other",
        "specific_reason": "detailed description",
        "original_text": "exact text from document",
        "affected_count": int,
        "severity": "low/medium/high/critical",
        "root_cause_analysis": "deep analysis of underlying cause",
        "recommended_intervention": "specific, actionable recommendation",
        "intervention_priority": "low/medium/high/urgent",
        "estimated_impact_if_resolved": "analysis of potential improvement"
      }}
    ],

    "mobilization_analysis": {{
      "asha_performance": {{
        "asha_name": "extracted or 'Unknown'",
        "home_visits_conducted": int or null,
        "beneficiaries_contacted": int or null,
        "advance_notice_given": "yes/no/partial/unknown",
        "notice_period_days": int or null,
        "communication_methods": ["home_visit/phone/sms/word_of_mouth/other"],
        "performance_rating": "poor/fair/good/excellent/unknown",
        "specific_failures": ["list of identified failures"],
        "specific_successes": ["list of identified successes"]
      }},
      "systemic_mobilization_issues": [
        {{
          "issue": "detailed description",
          "frequency": "isolated/recurring/systemic",
          "impact": "analysis",
          "intervention_needed": "specific recommendation"
        }}
      ]
    }},

    "demographic_insights": {{
      "age_distribution": "analysis if data available",
      "social_patterns": "any social/cultural patterns noticed",
      "economic_factors": "economic barriers or patterns",
      "seasonal_factors": "any seasonal/timing patterns"
    }}
  }},

  "clinical_services_delivered": {{
    "staff_present": [
      {{
        "designation": "Medical Officer/Nurse/ANM/Lab Tech/Pharmacist/Other",
        "name": "if mentioned",
        "present": true/false,
        "role_fulfillment": "how well they performed their role"
      }}
    ],
    "staff_adequacy": "severely_understaffed/understaffed/adequate/well_staffed/unknown",

    "physical_examination": {{
      "conducted": true/false/null,
      "beneficiaries_examined": int or null,
      "parameters_checked": ["height/weight/bmi/bp/pulse/temperature/pallor/thyroid/breast/abdominal/other"],
      "quality_of_examination": "cursory/standard/thorough/unknown",
      "abnormalities_detected": [
        {{
          "finding": "description",
          "severity": "low/medium/high",
          "action_taken": "what was done"
        }}
      ]
    }},

    "counselling_provided": {{
      "nutrition": {{"provided": true/false/null, "quality": "poor/fair/good/excellent/unknown"}},
      "exercise": {{"provided": true/false/null, "quality": "poor/fair/good/excellent/unknown"}},
      "family_planning": {{"provided": true/false/null, "quality": "poor/fair/good/excellent/unknown"}},
      "birth_spacing": {{"provided": true/false/null, "quality": "poor/fair/good/excellent/unknown"}},
      "anemia_prevention": {{"provided": true/false/null, "quality": "poor/fair/good/excellent/unknown"}},
      "hygiene": {{"provided": true/false/null, "quality": "poor/fair/good/excellent/unknown"}},
      "danger_signs": {{"provided": true/false/null, "quality": "poor/fair/good/excellent/unknown"}},
      "mental_health": {{"provided": true/false/null, "quality": "poor/fair/good/excellent/unknown"}},
      "overall_counselling_quality": "poor/fair/good/excellent/unknown",
      "gaps_in_counselling": ["list of topics not covered that should have been"]
    }},

    "laboratory_services": {{
      "tests_conducted": [
        {{
          "test_name": "hemoglobin/blood_sugar/thyroid/urine/hiv/hbsag/vdrl/blood_group/other",
          "samples_collected": int,
          "collection_quality": "poor/fair/good/excellent/unknown",
          "timing": "appropriate/delayed/unknown"
        }}
      ],
      "sample_handling": {{
        "storage_method": "refrigerated/room_temp/ice_box/unknown",
        "cold_chain_maintained": true/false/null,
        "storage_duration_hours": float or null,
        "transportation_method": "ambulance/cold_box/other/unknown",
        "critical_violations": ["list of violations"],
        "sample_integrity_risk": "none/low/medium/high/critical"
      }},
      "results_and_follow_up": {{
        "results_received": true/false/null,
        "turnaround_time_days": float or null,
        "results_shared_with_beneficiaries": true/false/null,
        "abnormal_results_action_taken": true/false/null,
        "follow_up_gaps": ["list of gaps in follow-up"]
      }},
      "lab_capacity_issues": [
        {{
          "issue_type": "equipment/staff/reagents/capacity/quality/other",
          "description": "detailed description",
          "impact": "impact on service delivery",
          "urgency": "low/medium/high/critical"
        }}
      ]
    }},

    "medications_distributed": [
      {{
        "medication": "IFA/calcium/folic_acid/deworming/other",
        "quantity_per_beneficiary": "amount",
        "beneficiaries_received": int,
        "counselling_on_usage": true/false/null,
        "stock_adequacy": "out_of_stock/low_stock/adequate/unknown"
      }}
    ],

    "service_quality_assessment": {{
      "overall_service_quality": 0-100,
      "strengths": ["list of what went well"],
      "weaknesses": ["list of gaps and problems"],
      "critical_failures": ["list of serious failures"],
      "best_practices_observed": ["list of exemplary practices"]
    }}
  }},

  "protocol_compliance_analysis": {{
    "due_list_preparation": {{
      "prepared": true/false/null,
      "quality": "poor/fair/good/excellent/unknown",
      "issues": "any problems with due list"
    }},
    "register_maintenance": {{
      "updated": true/false/null,
      "quality": "poor/fair/good/excellent/unknown",
      "gaps": "any gaps in documentation"
    }},
    "iec_materials": {{
      "displayed": true/false/null,
      "quality": "poor/fair/good/excellent/unknown",
      "effectiveness": "analysis of beneficiary awareness"
    }},
    "reporting_timeliness": {{
      "timely": true/false/null,
      "delay_if_any": "description of delay"
    }},

    "protocol_deviations": [
      {{
        "protocol": "which protocol was deviated from",
        "deviation": "what was done differently",
        "reason": "why it happened",
        "justification": "was it justified?",
        "impact": "what was the impact",
        "severity": "minor/moderate/major/critical",
        "corrective_action_needed": "what should be done"
      }}
    ],

    "compliance_score": 0-100,
    "compliance_rating": "non_compliant/partially_compliant/largely_compliant/fully_compliant"
  }},

  "risk_assessment": {{
    "immediate_risks": [
      {{
        "risk_type": "patient_safety/quality_of_care/legal/reputational/operational/other",
        "description": "detailed description",
        "likelihood": "low/medium/high/certain",
        "impact": "low/medium/high/severe",
        "risk_level": "low/medium/high/critical",
        "evidence": "what indicates this risk",
        "immediate_action_required": "what must be done now",
        "responsible_party": "who should address this",
        "timeline": "how urgently"
      }}
    ],
    "systemic_risks": [
      {{
        "risk_category": "category",
        "description": "systemic issue identified",
        "trend": "improving/stable/worsening/unknown",
        "broader_implications": "impact beyond this facility",
        "strategic_intervention_needed": "recommendation"
      }}
    ],
    "overall_risk_level": "low/medium/high/critical"
  }},

  "qualitative_insights": {{
    "field_worker_observations": {{
      "direct_quotes": ["exact quotes from the report"],
      "sentiment": "frustrated/neutral/satisfied/enthusiastic/mixed",
      "morale_indicators": "assessment of staff morale",
      "implicit_messages": "what is being communicated between the lines"
    }},

    "beneficiary_experience_indicators": {{
      "satisfaction_signals": ["any indications of beneficiary satisfaction"],
      "dissatisfaction_signals": ["any indications of problems"],
      "trust_in_system": "low/medium/high/unknown",
      "accessibility_perception": "analysis"
    }},

    "community_context": {{
      "cultural_factors": ["cultural elements affecting service delivery"],
      "social_dynamics": ["social patterns observed"],
      "economic_context": ["economic factors at play"],
      "seasonal_considerations": ["any seasonal factors"]
    }},

    "emerging_patterns": [
      {{
        "pattern": "description of pattern",
        "significance": "why this matters",
        "evidence": "what supports this observation",
        "recommendation": "what should be done"
      }}
    ]
  }},

  "intelligent_analysis": {{
    "root_cause_analysis": [
      {{
        "problem": "main problem identified",
        "surface_causes": ["immediate causes"],
        "underlying_causes": ["deeper systemic causes"],
        "contributing_factors": ["other factors"],
        "evidence_chain": "how we know this",
        "confidence_level": "low/medium/high",
        "verification_needed": "what additional data would confirm this"
      }}
    ],

    "performance_benchmarking": {{
      "strengths": ["areas of excellence"],
      "weaknesses": ["areas needing improvement"],
      "comparison_to_standards": "how this compares to protocols",
      "improvement_potential": "realistic improvement possibilities",
      "quick_wins": ["easy improvements with high impact"],
      "long_term_investments": ["structural improvements needed"]
    }},

    "predictive_insights": [
      {{
        "trend": "identified trend",
        "trajectory": "improving/declining/stable/volatile",
        "projection": "where this is heading",
        "intervention_window": "when to act",
        "prevention_opportunity": "how to prevent problems"
      }}
    ],

    "comparative_analysis": {{
      "versus_previous_reports": "comparison if historical context known",
      "versus_peer_facilities": "comparison if peer data available",
      "versus_targets": "comparison to government targets"
    }}
  }},

  "actionable_recommendations": {{
    "immediate_actions": [
      {{
        "priority": 1-10,
        "action": "specific action to take",
        "rationale": "why this is needed",
        "responsible_party": "who should do this",
        "timeline": "when it should be done",
        "resources_required": ["what resources are needed"],
        "expected_outcome": "what will improve",
        "success_metrics": ["how to measure success"],
        "implementation_difficulty": "easy/medium/hard/very_hard"
      }}
    ],

    "short_term_improvements": [
      {{
        "timeframe": "1-3 months",
        "improvement": "what to improve",
        "approach": "how to improve it",
        "expected_impact": "projected benefit",
        "dependencies": ["what needs to happen first"]
      }}
    ],

    "strategic_initiatives": [
      {{
        "timeframe": "3-12 months",
        "initiative": "strategic change needed",
        "business_case": "why invest in this",
        "stakeholders": ["who needs to be involved"],
        "estimated_roi": "return on investment analysis"
      }}
    ],

    "capacity_building_needs": [
      {{
        "area": "what capability needs building",
        "current_level": "assessment of current state",
        "target_level": "desired state",
        "training_needed": "specific training requirements",
        "mentoring_needed": "ongoing support needed",
        "infrastructure_needed": "resources required"
      }}
    ]
  }},

  "executive_summary": {{
    "one_sentence_summary": "single sentence capturing essence",
    "key_findings": ["top 5-7 findings"],
    "critical_issues": ["issues requiring immediate attention"],
    "positive_highlights": ["things that went well"],
    "overall_assessment": "comprehensive 2-3 paragraph assessment",
    "bottom_line": "clear conclusion about facility performance"
  }},

  "meta_analysis": {{
    "data_confidence": {{
      "overall_confidence": 0.0-1.0,
      "high_confidence_areas": ["areas where data is reliable"],
      "low_confidence_areas": ["areas with uncertain data"],
      "assumptions_made": ["key assumptions in this analysis"],
      "validation_needed": ["what should be verified"]
    }},
    "analysis_limitations": ["limitations of this analysis"],
    "additional_data_needed": ["what additional information would help"],
    "follow_up_questions": ["questions for facility staff"]
  }}
}}

CRITICAL INSTRUCTIONS:
1. Extract EVERYTHING possible from the text, even if you have to infer
2. Be EXTREMELY detailed and granular - this needs to impress government officials
3. Provide deep root cause analysis, not just surface observations
4. Make specific, actionable recommendations with clear priorities
5. Connect dots between different issues to show systemic patterns
6. Quantify whenever possible (scores, percentages, counts)
7. Be honest about confidence levels and data quality
8. Highlight both problems AND successes
9. Think like a senior public health official reviewing this for policy decisions

Return ONLY the JSON object, no other text."""

        print("🤖 Analyzing document with AI (this may take 30-60 seconds)...")

        response = self.client.messages.create(
            model=self.model,
            max_tokens=16000,
            temperature=0.3,  # Lower temperature for more consistent extraction
            messages=[
                {
                    "role": "user",
                    "content": analysis_prompt
                }
            ]
        )

        # Extract JSON from response
        response_text = response.content[0].text

        # Try to parse JSON
        try:
            # Sometimes Claude wraps JSON in markdown code blocks
            if "```json" in response_text:
                json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
                if json_match:
                    response_text = json_match.group(1)
            elif "```" in response_text:
                json_match = re.search(r'```\s*(.*?)\s*```', response_text, re.DOTALL)
                if json_match:
                    response_text = json_match.group(1)

            analysis = json.loads(response_text)
            return analysis

        except json.JSONDecodeError as e:
            print(f"ERROR: Failed to parse AI response as JSON: {e}")
            print(f"Response preview: {response_text[:500]}...")
            raise

    def generate_report(self, analysis: Dict[str, Any], output_path: str):
        """Generate comprehensive markdown report."""

        report = f"""# 🏥 GOVERNMENT HEALTH DATA ANALYSIS REPORT
## Preconception & Maternal Health Clinic (PPC) Assessment

---

**Report Generated**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
**Analysis System**: Government Health Data Automation Platform (Post-POC MVP)
**Analysis Model**: Claude Sonnet 4.5 (Anthropic)

---

# 📊 EXECUTIVE SUMMARY

## {analysis['executive_summary']['one_sentence_summary']}

{analysis['executive_summary']['overall_assessment']}

### Bottom Line
{analysis['executive_summary']['bottom_line']}

---

## 🎯 KEY FINDINGS

"""

        for i, finding in enumerate(analysis['executive_summary']['key_findings'], 1):
            report += f"{i}. **{finding}**\n"

        report += "\n---\n\n"

        # Critical Issues
        if analysis['executive_summary']['critical_issues']:
            report += "## 🚨 CRITICAL ISSUES REQUIRING IMMEDIATE ATTENTION\n\n"
            for i, issue in enumerate(analysis['executive_summary']['critical_issues'], 1):
                report += f"{i}. ⚠️ **{issue}**\n"
            report += "\n---\n\n"

        # Positive Highlights
        if analysis['executive_summary']['positive_highlights']:
            report += "## ✅ POSITIVE HIGHLIGHTS\n\n"
            for i, highlight in enumerate(analysis['executive_summary']['positive_highlights'], 1):
                report += f"{i}. ✨ {highlight}\n"
            report += "\n---\n\n"

        # Document Quality Assessment
        report += f"""# 📄 DOCUMENT QUALITY ASSESSMENT

**Readability Score**: {analysis['document_metadata']['document_quality_assessment']['readability_score']}/100
**Completeness Score**: {analysis['document_metadata']['document_quality_assessment']['completeness_score']}/100
**Data Clarity**: {analysis['document_metadata']['document_quality_assessment']['data_clarity'].upper()}
**Language Detected**: {analysis['document_metadata']['document_quality_assessment']['language_detected'].title()}

### Special Challenges in Data Extraction
"""

        for challenge in analysis['document_metadata']['document_quality_assessment']['special_challenges']:
            report += f"- {challenge}\n"

        report += "\n---\n\n"

        # Facility Information
        fac = analysis['facility_information']
        report += f"""# 🏢 FACILITY INFORMATION

**Facility Name**: {fac['facility_name']}
**Type**: {fac['facility_type']}
**Location**: {fac['location']['block']}, {fac['location']['district']}, {fac['location']['state']}

**Data Confidence**:
- Facility Name: {fac['confidence_scores']['facility_name']:.0%}
- Location: {fac['confidence_scores']['location']:.0%}

**Extraction Notes**: {fac['extraction_notes']}

---

"""

        # Clinic Details
        clinic = analysis['clinic_details']
        report += f"""# 📅 CLINIC DETAILS

**Date**: {clinic['clinic_date'] or 'Not specified'}
**Day**: {clinic['day_of_week'] or 'Unknown'}
**Type**: {clinic['clinic_type']}

**Timing**:
- Start: {clinic['timing']['start_time'] or 'Not specified'}
- End: {clinic['timing']['end_time'] or 'Not specified'}
- Duration: {clinic['timing']['duration_hours'] or 'Unknown'} hours

**Reporting Staff**:
- Name: {clinic['reporting_staff']['name']}
- Designation: {clinic['reporting_staff']['designation']}
- Contact: {clinic['reporting_staff']['contact'] or 'Not provided'}

---

"""

        # Infrastructure Assessment
        infra = analysis['infrastructure_assessment']
        report += f"""# 🏗️ INFRASTRUCTURE ASSESSMENT

**Overall Infrastructure Score**: {infra['overall_infrastructure_score']}/100

## Physical Space
- **Dedicated Space Available**: {infra['dedicated_space']['available']}
- **Type**: {infra['dedicated_space']['type'].replace('_', ' ').title()}
- **Privacy Ensured**: {infra['dedicated_space']['privacy_ensured']}

### Issues Identified
"""

        for issue in infra['dedicated_space']['issues_identified']:
            report += f"- {issue}\n"

        report += "\n## Waiting Area\n"
        report += f"- **Available**: {infra['waiting_area']['available']}\n"
        report += f"- **Adequacy**: {infra['waiting_area']['adequacy'].replace('_', ' ').title()}\n\n"

        report += "## Signage and Awareness\n"
        report += f"- **PPC Board Displayed**: {infra['signage_and_awareness']['ppc_board_displayed']}\n"
        report += f"- **Beneficiary Awareness**: {infra['signage_and_awareness']['beneficiary_awareness'].title()}\n\n"

        if infra['critical_gaps']:
            report += "## 🔴 CRITICAL INFRASTRUCTURE GAPS\n\n"
            for gap in infra['critical_gaps']:
                report += f"""### {gap['gap_type'].upper()} - Severity: {gap['severity'].upper()}

**Description**: {gap['description']}

**Evidence**: "{gap['evidence_from_text']}"

**Potential Impact**: {gap['potential_impact']}

---

"""

        # Beneficiary Attendance Analysis
        att = analysis['beneficiary_attendance_analysis']
        report += f"""# 👥 BENEFICIARY ATTENDANCE ANALYSIS

**Expected**: {att['expected_beneficiaries']}
**Actual**: {att['actual_attendance']}
**Attendance Rate**: {att['attendance_rate']:.1%} - **{att['attendance_quality'].replace('_', ' ').upper()}**

"""

        if att['individual_beneficiaries']:
            report += "## Individual Beneficiaries\n\n"
            for ben in att['individual_beneficiaries']:
                report += f"""### {ben['beneficiary_identifier']}
- **Attended**: {ben['attended']}
- **Age**: {ben['age']}
- **Status**: {ben['gestational_status'].title()}
"""
                if not ben['attended'] and ben['reason_for_absence']:
                    report += f"- **Reason for Absence**: {ben['reason_for_absence']}\n"
                if ben['family_situation']:
                    report += f"- **Family Context**: {ben['family_situation']}\n"
                report += "\n"

        if att['attendance_barriers_detected']:
            report += "## 🚧 ATTENDANCE BARRIERS - DEEP ANALYSIS\n\n"
            for i, barrier in enumerate(att['attendance_barriers_detected'], 1):
                report += f"""### Barrier #{i}: {barrier['barrier_type'].upper()} - Severity: {barrier['severity'].upper()}

**Specific Reason**: {barrier['specific_reason']}

**Original Text**: "{barrier['original_text']}"

**Affected Count**: {barrier['affected_count']} beneficiaries

#### 🔍 Root Cause Analysis
{barrier['root_cause_analysis']}

#### 💡 Recommended Intervention (Priority: {barrier['intervention_priority'].upper()})
{barrier['recommended_intervention']}

#### 📈 Estimated Impact if Resolved
{barrier['estimated_impact_if_resolved']}

---

"""

        # ASHA Performance
        asha = att['mobilization_analysis']['asha_performance']
        report += f"""## 👩‍⚕️ ASHA PERFORMANCE ANALYSIS

**ASHA Name**: {asha['asha_name']}
**Performance Rating**: {asha['performance_rating'].upper()}

### Mobilization Metrics
- **Home Visits Conducted**: {asha['home_visits_conducted']}
- **Beneficiaries Contacted**: {asha['beneficiaries_contacted']}
- **Advance Notice Given**: {asha['advance_notice_given'].title()}
- **Notice Period**: {asha['notice_period_days']} days
- **Communication Methods**: {', '.join([m.replace('_', ' ').title() for m in asha['communication_methods']])}

### Specific Failures
"""

        for failure in asha['specific_failures']:
            report += f"- ❌ {failure}\n"

        report += "\n### Specific Successes\n"
        for success in asha['specific_successes']:
            report += f"- ✅ {success}\n"

        report += "\n"

        if att['mobilization_analysis']['systemic_mobilization_issues']:
            report += "## ⚙️ SYSTEMIC MOBILIZATION ISSUES\n\n"
            for issue in att['mobilization_analysis']['systemic_mobilization_issues']:
                report += f"""### {issue['issue']}
- **Frequency**: {issue['frequency'].title()}
- **Impact**: {issue['impact']}
- **Intervention Needed**: {issue['intervention_needed']}

"""

        # Demographic Insights
        demo = att['demographic_insights']
        report += f"""## 📊 DEMOGRAPHIC INSIGHTS

**Age Distribution**: {demo['age_distribution']}

**Social Patterns**: {demo['social_patterns']}

**Economic Factors**: {demo['economic_factors']}

**Seasonal Factors**: {demo['seasonal_factors']}

---

"""

        # Clinical Services
        clin = analysis['clinical_services_delivered']
        report += f"""# 💉 CLINICAL SERVICES DELIVERED

**Staff Adequacy**: {clin['staff_adequacy'].replace('_', ' ').upper()}

## Staff Present

"""

        for staff in clin['staff_present']:
            report += f"""### {staff['designation']}
- **Name**: {staff.get('name', 'Not specified')}
- **Present**: {staff['present']}
- **Role Fulfillment**: {staff['role_fulfillment']}

"""

        # Physical Examination
        exam = clin['physical_examination']
        report += f"""## 🔬 Physical Examination

- **Conducted**: {exam['conducted']}
- **Beneficiaries Examined**: {exam['beneficiaries_examined']}
- **Quality**: {exam['quality_of_examination'].replace('_', ' ').title()}
- **Parameters Checked**: {', '.join([p.replace('_', ' ').title() for p in exam['parameters_checked']])}

"""

        if exam['abnormalities_detected']:
            report += "### Abnormalities Detected\n\n"
            for abn in exam['abnormalities_detected']:
                report += f"- **{abn['finding']}** (Severity: {abn['severity'].title()}) - Action: {abn['action_taken']}\n"
            report += "\n"

        # Counselling
        report += "## 🗣️ COUNSELLING PROVIDED\n\n"
        report += f"**Overall Quality**: {clin['counselling_provided']['overall_counselling_quality'].replace('_', ' ').upper()}\n\n"

        report += "| Topic | Provided | Quality |\n"
        report += "|-------|----------|----------|\n"

        counselling_topics = [
            'nutrition', 'exercise', 'family_planning', 'birth_spacing',
            'anemia_prevention', 'hygiene', 'danger_signs', 'mental_health'
        ]

        for topic in counselling_topics:
            if topic in clin['counselling_provided']:
                data = clin['counselling_provided'][topic]
                report += f"| {topic.replace('_', ' ').title()} | {data['provided']} | {data['quality'].replace('_', ' ').title()} |\n"

        report += "\n"

        if clin['counselling_provided']['gaps_in_counselling']:
            report += "### Gaps in Counselling\n"
            for gap in clin['counselling_provided']['gaps_in_counselling']:
                report += f"- ⚠️ {gap}\n"
            report += "\n"

        # Laboratory Services
        lab = clin['laboratory_services']
        report += "## 🧪 LABORATORY SERVICES\n\n"

        if lab['tests_conducted']:
            report += "### Tests Conducted\n\n"
            report += "| Test | Samples | Quality | Timing |\n"
            report += "|------|---------|---------|--------|\n"
            for test in lab['tests_conducted']:
                report += f"| {test['test_name'].replace('_', ' ').title()} | {test['samples_collected']} | {test['collection_quality'].title()} | {test['timing'].title()} |\n"
            report += "\n"

        # Sample Handling
        handling = lab['sample_handling']
        report += f"""### Sample Handling & Storage

- **Storage Method**: {handling['storage_method'].replace('_', ' ').title()}
- **Cold Chain Maintained**: {handling['cold_chain_maintained']}
- **Storage Duration**: {handling['storage_duration_hours']} hours
- **Transportation**: {handling['transportation_method'].replace('_', ' ').title()}
- **Sample Integrity Risk**: {handling['sample_integrity_risk'].upper()}

"""

        if handling['critical_violations']:
            report += "#### 🚨 CRITICAL VIOLATIONS\n"
            for violation in handling['critical_violations']:
                report += f"- ⛔ {violation}\n"
            report += "\n"

        # Results and Follow-up
        results = lab['results_and_follow_up']
        report += f"""### Results & Follow-up

- **Results Received**: {results['results_received']}
- **Turnaround Time**: {results['turnaround_time_days']} days
- **Results Shared with Beneficiaries**: {results['results_shared_with_beneficiaries']}
- **Abnormal Results Action Taken**: {results['abnormal_results_action_taken']}

"""

        if results['follow_up_gaps']:
            report += "#### Follow-up Gaps\n"
            for gap in results['follow_up_gaps']:
                report += f"- {gap}\n"
            report += "\n"

        if lab['lab_capacity_issues']:
            report += "### ⚠️ Lab Capacity Issues\n\n"
            for issue in lab['lab_capacity_issues']:
                report += f"""#### {issue['issue_type'].upper()} - Urgency: {issue['urgency'].upper()}

**Description**: {issue['description']}

**Impact**: {issue['impact']}

---

"""

        # Medications
        if clin['medications_distributed']:
            report += "## 💊 MEDICATIONS DISTRIBUTED\n\n"
            report += "| Medication | Quantity | Beneficiaries | Counselling | Stock Status |\n"
            report += "|------------|----------|---------------|-------------|---------------|\n"
            for med in clin['medications_distributed']:
                report += f"| {med['medication'].replace('_', ' ').upper()} | {med['quantity_per_beneficiary']} | {med['beneficiaries_received']} | {med['counselling_on_usage']} | {med['stock_adequacy'].replace('_', ' ').title()} |\n"
            report += "\n"

        # Service Quality
        quality = clin['service_quality_assessment']
        report += f"""## 📈 SERVICE QUALITY ASSESSMENT

**Overall Score**: {quality['overall_service_quality']}/100

### Strengths
"""

        for strength in quality['strengths']:
            report += f"- ✅ {strength}\n"

        report += "\n### Weaknesses\n"
        for weakness in quality['weaknesses']:
            report += f"- ⚠️ {weakness}\n"

        if quality['critical_failures']:
            report += "\n### 🔴 Critical Failures\n"
            for failure in quality['critical_failures']:
                report += f"- ⛔ {failure}\n"

        if quality['best_practices_observed']:
            report += "\n### 🌟 Best Practices Observed\n"
            for practice in quality['best_practices_observed']:
                report += f"- ⭐ {practice}\n"

        report += "\n---\n\n"

        # Protocol Compliance
        comp = analysis['protocol_compliance_analysis']
        report += f"""# 📋 PROTOCOL COMPLIANCE ANALYSIS

**Compliance Score**: {comp['compliance_score']}/100
**Rating**: {comp['compliance_rating'].replace('_', ' ').upper()}

## Compliance Checklist

### Due List Preparation
- **Prepared**: {comp['due_list_preparation']['prepared']}
- **Quality**: {comp['due_list_preparation']['quality'].title()}
- **Issues**: {comp['due_list_preparation']['issues']}

### Register Maintenance
- **Updated**: {comp['register_maintenance']['updated']}
- **Quality**: {comp['register_maintenance']['quality'].title()}
- **Gaps**: {comp['register_maintenance']['gaps']}

### IEC Materials
- **Displayed**: {comp['iec_materials']['displayed']}
- **Quality**: {comp['iec_materials']['quality'].title()}
- **Effectiveness**: {comp['iec_materials']['effectiveness']}

### Reporting Timeliness
- **Timely**: {comp['reporting_timeliness']['timely']}
- **Delay**: {comp['reporting_timeliness']['delay_if_any']}

"""

        if comp['protocol_deviations']:
            report += "## 🚨 PROTOCOL DEVIATIONS\n\n"
            for dev in comp['protocol_deviations']:
                report += f"""### {dev['protocol']}

**Deviation**: {dev['deviation']}

**Reason**: {dev['reason']}

**Justification**: {dev['justification']}

**Impact**: {dev['impact']}

**Severity**: {dev['severity'].upper()}

**Corrective Action**: {dev['corrective_action_needed']}

---

"""

        # Risk Assessment
        risk = analysis['risk_assessment']
        report += f"""# ⚠️ RISK ASSESSMENT

**Overall Risk Level**: {risk['overall_risk_level'].upper()}

"""

        if risk['immediate_risks']:
            report += "## 🚨 IMMEDIATE RISKS\n\n"
            for r in risk['immediate_risks']:
                report += f"""### {r['risk_type'].upper()} - Risk Level: {r['risk_level'].upper()}

**Description**: {r['description']}

**Likelihood**: {r['likelihood'].title()} | **Impact**: {r['impact'].title()}

**Evidence**: {r['evidence']}

**Immediate Action Required**: {r['immediate_action_required']}

**Responsible Party**: {r['responsible_party']}

**Timeline**: {r['timeline']}

---

"""

        if risk['systemic_risks']:
            report += "## ⚙️ SYSTEMIC RISKS\n\n"
            for r in risk['systemic_risks']:
                report += f"""### {r['risk_category'].upper()}

**Description**: {r['description']}

**Trend**: {r['trend'].title()}

**Broader Implications**: {r['broader_implications']}

**Strategic Intervention**: {r['strategic_intervention_needed']}

---

"""

        # Qualitative Insights
        qual = analysis['qualitative_insights']
        report += "# 🔍 QUALITATIVE INSIGHTS\n\n"

        report += "## Field Worker Observations\n\n"

        if qual['field_worker_observations']['direct_quotes']:
            report += "### Direct Quotes\n"
            for quote in qual['field_worker_observations']['direct_quotes']:
                report += f'> "{quote}"\n\n'

        report += f"""**Sentiment**: {qual['field_worker_observations']['sentiment'].title()}

**Morale Indicators**: {qual['field_worker_observations']['morale_indicators']}

**Implicit Messages**: {qual['field_worker_observations']['implicit_messages']}

"""

        report += "## Beneficiary Experience Indicators\n\n"

        if qual['beneficiary_experience_indicators']['satisfaction_signals']:
            report += "### Satisfaction Signals\n"
            for signal in qual['beneficiary_experience_indicators']['satisfaction_signals']:
                report += f"- ✅ {signal}\n"
            report += "\n"

        if qual['beneficiary_experience_indicators']['dissatisfaction_signals']:
            report += "### Dissatisfaction Signals\n"
            for signal in qual['beneficiary_experience_indicators']['dissatisfaction_signals']:
                report += f"- ❌ {signal}\n"
            report += "\n"

        report += f"""**Trust in System**: {qual['beneficiary_experience_indicators']['trust_in_system'].title()}

**Accessibility Perception**: {qual['beneficiary_experience_indicators']['accessibility_perception']}

"""

        report += "## Community Context\n\n"

        if qual['community_context']['cultural_factors']:
            report += "### Cultural Factors\n"
            for factor in qual['community_context']['cultural_factors']:
                report += f"- {factor}\n"
            report += "\n"

        if qual['community_context']['social_dynamics']:
            report += "### Social Dynamics\n"
            for dynamic in qual['community_context']['social_dynamics']:
                report += f"- {dynamic}\n"
            report += "\n"

        if qual['community_context']['economic_context']:
            report += "### Economic Context\n"
            for context in qual['community_context']['economic_context']:
                report += f"- {context}\n"
            report += "\n"

        if qual['community_context']['seasonal_considerations']:
            report += "### Seasonal Considerations\n"
            for consideration in qual['community_context']['seasonal_considerations']:
                report += f"- {consideration}\n"
            report += "\n"

        if qual['emerging_patterns']:
            report += "## 📈 EMERGING PATTERNS\n\n"
            for pattern in qual['emerging_patterns']:
                report += f"""### {pattern['pattern']}

**Significance**: {pattern['significance']}

**Evidence**: {pattern['evidence']}

**Recommendation**: {pattern['recommendation']}

---

"""

        # Intelligent Analysis
        intel = analysis['intelligent_analysis']
        report += "# 🧠 INTELLIGENT ANALYSIS\n\n"

        if intel['root_cause_analysis']:
            report += "## 🔬 ROOT CAUSE ANALYSIS\n\n"
            for rca in intel['root_cause_analysis']:
                report += f"""### Problem: {rca['problem']}

#### Surface Causes
"""
                for cause in rca['surface_causes']:
                    report += f"- {cause}\n"

                report += "\n#### Underlying Causes\n"
                for cause in rca['underlying_causes']:
                    report += f"- {cause}\n"

                report += "\n#### Contributing Factors\n"
                for factor in rca['contributing_factors']:
                    report += f"- {factor}\n"

                report += f"""
**Evidence Chain**: {rca['evidence_chain']}

**Confidence Level**: {rca['confidence_level'].upper()}

**Verification Needed**: {rca['verification_needed']}

---

"""

        # Performance Benchmarking
        bench = intel['performance_benchmarking']
        report += f"""## 📊 PERFORMANCE BENCHMARKING

### Strengths
"""

        for strength in bench['strengths']:
            report += f"- ✅ {strength}\n"

        report += "\n### Weaknesses\n"
        for weakness in bench['weaknesses']:
            report += f"- ⚠️ {weakness}\n"

        report += f"""
**Comparison to Standards**: {bench['comparison_to_standards']}

**Improvement Potential**: {bench['improvement_potential']}

### Quick Wins (Easy improvements with high impact)
"""

        for win in bench['quick_wins']:
            report += f"- 🎯 {win}\n"

        report += "\n### Long-term Investments\n"
        for inv in bench['long_term_investments']:
            report += f"- 📈 {inv}\n"

        report += "\n"

        # Predictive Insights
        if intel['predictive_insights']:
            report += "## 🔮 PREDICTIVE INSIGHTS\n\n"
            for pred in intel['predictive_insights']:
                report += f"""### Trend: {pred['trend']}

**Trajectory**: {pred['trajectory'].title()}

**Projection**: {pred['projection']}

**Intervention Window**: {pred['intervention_window']}

**Prevention Opportunity**: {pred['prevention_opportunity']}

---

"""

        # Comparative Analysis
        comp_analysis = intel['comparative_analysis']
        report += f"""## 📐 COMPARATIVE ANALYSIS

**Versus Previous Reports**: {comp_analysis['versus_previous_reports']}

**Versus Peer Facilities**: {comp_analysis['versus_peer_facilities']}

**Versus Targets**: {comp_analysis['versus_targets']}

---

"""

        # Actionable Recommendations
        rec = analysis['actionable_recommendations']
        report += "# 💡 ACTIONABLE RECOMMENDATIONS\n\n"

        if rec['immediate_actions']:
            report += "## 🚨 IMMEDIATE ACTIONS (Priority Ordered)\n\n"
            for action in sorted(rec['immediate_actions'], key=lambda x: x['priority']):
                report += f"""### Priority {action['priority']}: {action['action']}

**Rationale**: {action['rationale']}

**Responsible Party**: {action['responsible_party']}

**Timeline**: {action['timeline']}

**Resources Required**:
"""
                for resource in action['resources_required']:
                    report += f"- {resource}\n"

                report += f"""
**Expected Outcome**: {action['expected_outcome']}

**Success Metrics**:
"""
                for metric in action['success_metrics']:
                    report += f"- {metric}\n"

                report += f"""
**Implementation Difficulty**: {action['implementation_difficulty'].replace('_', ' ').title()}

---

"""

        if rec['short_term_improvements']:
            report += "## 📅 SHORT-TERM IMPROVEMENTS (1-3 months)\n\n"
            for imp in rec['short_term_improvements']:
                report += f"""### {imp['improvement']}

**Timeframe**: {imp['timeframe']}

**Approach**: {imp['approach']}

**Expected Impact**: {imp['expected_impact']}

**Dependencies**:
"""
                for dep in imp['dependencies']:
                    report += f"- {dep}\n"

                report += "\n---\n\n"

        if rec['strategic_initiatives']:
            report += "## 🎯 STRATEGIC INITIATIVES (3-12 months)\n\n"
            for init in rec['strategic_initiatives']:
                report += f"""### {init['initiative']}

**Timeframe**: {init['timeframe']}

**Business Case**: {init['business_case']}

**Stakeholders**:
"""
                for stakeholder in init['stakeholders']:
                    report += f"- {stakeholder}\n"

                report += f"""
**Estimated ROI**: {init['estimated_roi']}

---

"""

        if rec['capacity_building_needs']:
            report += "## 👨‍🎓 CAPACITY BUILDING NEEDS\n\n"
            for need in rec['capacity_building_needs']:
                report += f"""### {need['area']}

**Current Level**: {need['current_level']}

**Target Level**: {need['target_level']}

**Training Needed**: {need['training_needed']}

**Mentoring Needed**: {need['mentoring_needed']}

**Infrastructure Needed**: {need['infrastructure_needed']}

---

"""

        # Meta Analysis
        meta = analysis['meta_analysis']
        report += f"""# 🔬 META-ANALYSIS & DATA QUALITY

## Data Confidence Assessment

**Overall Confidence**: {meta['data_confidence']['overall_confidence']:.0%}

### High Confidence Areas
"""

        for area in meta['data_confidence']['high_confidence_areas']:
            report += f"- ✅ {area}\n"

        report += "\n### Low Confidence Areas\n"
        for area in meta['data_confidence']['low_confidence_areas']:
            report += f"- ⚠️ {area}\n"

        report += "\n### Assumptions Made\n"
        for assumption in meta['data_confidence']['assumptions_made']:
            report += f"- {assumption}\n"

        report += "\n### Validation Needed\n"
        for validation in meta['data_confidence']['validation_needed']:
            report += f"- {validation}\n"

        report += "\n## Analysis Limitations\n\n"
        for limitation in meta['analysis_limitations']:
            report += f"- {limitation}\n"

        report += "\n## Additional Data Needed\n\n"
        for data in meta['additional_data_needed']:
            report += f"- {data}\n"

        report += "\n## Follow-up Questions for Facility Staff\n\n"
        for question in meta['follow_up_questions']:
            report += f"- {question}\n"

        report += f"""

---

# 📝 CONCLUSION

This comprehensive analysis demonstrates the capability of AI-powered automation to extract deep, actionable insights from field reports, regardless of data quality or format. The system provides:

1. **Structured Data Extraction**: Converting messy, multilingual reports into standardized formats
2. **Deep Root Cause Analysis**: Going beyond surface observations to identify systemic issues
3. **Risk Assessment**: Identifying and prioritizing immediate and systemic risks
4. **Actionable Recommendations**: Providing specific, prioritized interventions with clear ownership
5. **Quality Metrics**: Quantifying performance across multiple dimensions
6. **Predictive Insights**: Identifying trends and intervention windows

This level of granular analysis enables data-driven decision-making at scale, replacing manual coordination while maintaining the human insight needed for effective public health management.

---

**Report End**
**Generated by**: Government Health Data Automation Platform
**Technology**: Claude Sonnet 4.5 (Anthropic AI)
**Analysis Date**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

---

*For questions or clarifications, please contact the analysis team.*
"""

        # Write report
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report)

        print(f"\n✅ Report generated: {output_path}")

        # Also save JSON
        json_path = output_path.replace('.md', '.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, indent=2, ensure_ascii=False)

        print(f"✅ JSON analysis saved: {json_path}")


def main():
    """Main entry point."""

    print("""
╔══════════════════════════════════════════════════════════════════════════════╗
║                                                                              ║
║        🏥 GOVERNMENT HEALTH DATA AUTOMATION - POST-POC MVP 🏥                ║
║                                                                              ║
║                    AI-Powered Granular Report Analysis                       ║
║                                                                              ║
╚══════════════════════════════════════════════════════════════════════════════╝
""")

    # Check arguments
    if len(sys.argv) != 2:
        print("Usage: python poc_analyzer.py <path_to_report.docx>")
        print("\nExample: python poc_analyzer.py data/samples/ppc_report.docx")
        sys.exit(1)

    input_file = sys.argv[1]

    # Check if file exists
    if not Path(input_file).exists():
        print(f"❌ ERROR: File not found: {input_file}")
        sys.exit(1)

    # Check API key
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("❌ ERROR: ANTHROPIC_API_KEY environment variable not set")
        print("\nGet your API key from: https://console.anthropic.com/")
        print("Then set it: export ANTHROPIC_API_KEY='your-key-here'")
        sys.exit(1)

    try:
        # Initialize analyzer
        print("🚀 Initializing AI-powered analyzer...")
        analyzer = GovernmentHealthAnalyzer()

        # Extract text
        print(f"📄 Extracting text from: {input_file}")
        raw_text = analyzer.extract_text(input_file)

        print(f"✅ Extracted {len(raw_text)} characters")
        print(f"📝 Text preview: {raw_text[:200]}...")

        # Analyze
        print("\n" + "="*80)
        analysis = analyzer.analyze_document(raw_text)

        # Generate report
        output_path = Path(input_file).stem + "_ANALYSIS_REPORT.md"
        print(f"\n📊 Generating comprehensive report...")
        analyzer.generate_report(analysis, output_path)

        print("\n" + "="*80)
        print("\n🎉 SUCCESS! Analysis complete.")
        print(f"\n📄 Reports generated:")
        print(f"   - Markdown: {output_path}")
        print(f"   - JSON: {output_path.replace('.md', '.json')}")

        print(f"\n📖 View report: open {output_path}")

        print("\n" + "="*80)
        print("\n✨ This analysis demonstrates:")
        print("   ✅ AI-powered extraction from unpredictable input")
        print("   ✅ Deep granular analysis worthy of government contracts")
        print("   ✅ Root cause analysis and predictive insights")
        print("   ✅ Actionable recommendations with clear priorities")
        print("   ✅ Comprehensive risk assessment")
        print("   ✅ Quality metrics and confidence scores")

        print("\n🚀 Ready for production deployment!\n")

    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
