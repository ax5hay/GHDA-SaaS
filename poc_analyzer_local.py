#!/usr/bin/env python3
"""
Government Health Data Automation - Post-POC MVP (LOCAL LM STUDIO VERSION)
Single-file analyzer using LOCAL AI models via LM Studio

This version uses your local LM Studio models instead of cloud APIs.

LM Studio Setup:
1. Start LM Studio
2. Load your preferred model (Qwen, Llama, Mistral, etc.)
3. Start the local server (default: http://192.168.56.1:1234)
4. Note the exact model name shown in LM Studio
5. Run this script with the model name

Usage:
    python poc_analyzer_local.py path/to/report.docx

    # Or set model name via environment variable:
    set LM_STUDIO_MODEL=qwen2.5-14b-instruct
    python poc_analyzer_local.py path/to/report.docx

Common model names:
    - qwen2.5 or qwen2.5-14b-instruct
    - llama-3.1-8b-instruct
    - mistral-7b-instruct
    - (check LM Studio for the exact name)

Requirements:
    pip install openai python-docx PyPDF2 python-dateutil
"""

import os
import re
import sys
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional

# Document processing imports
try:
    import docx
    from PyPDF2 import PdfReader
except ImportError:
    print("ERROR: Required packages not installed.")
    print("Install with: pip install openai python-docx PyPDF2 python-dateutil")
    sys.exit(1)

# OpenAI-compatible client for LM Studio
try:
    from openai import OpenAI
except ImportError:
    print("ERROR: OpenAI package not installed.")
    print("Install with: pip install openai")
    sys.exit(1)


class LocalGovernmentHealthAnalyzer:
    """
    Post-POC MVP: Local AI-powered health report analyzer using LM Studio
    """

    def __init__(self, base_url: str = "http://192.168.56.1:1234/v1", model_name: str = None):
        """Initialize analyzer with local LM Studio server."""
        self.base_url = base_url

        # Allow model name to be configured via env var or parameter
        self.model_name = model_name or os.environ.get("LM_STUDIO_MODEL", "phi-4-reasoning-plus")

        # Create OpenAI-compatible client pointing to LM Studio
        self.client = OpenAI(
            base_url=self.base_url,
            api_key="lm-studio"  # LM Studio doesn't need real API key
        )

        print(f"✅ Connected to local LM Studio at: {self.base_url}")
        print(f"🤖 Using model: {self.model_name}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                text_parts.append(" | ".join(row_data))

        return "\n".join(text_parts)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)

        return "\n".join(text_parts)

    def extract_text(self, file_path: str) -> str:
        """Extract text from document (DOCX or PDF)."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension == ".docx":
            return self.extract_text_from_docx(str(file_path))
        elif extension == ".pdf":
            return self.extract_text_from_pdf(str(file_path))
        else:
            return file_path.read_text(encoding="utf-8", errors="ignore")

    def analyze_document(self, raw_text: str) -> Dict[str, Any]:
        """
        AI-powered document analysis using local LM Studio model.
        """

        # Shorter, more focused prompt for local models
        analysis_prompt = f"""You are a government health data analyst for Maternal Health Clinics in India.

Analyze this field report and extract ALL information into structured JSON.

REPORT:
{raw_text}

---

Extract this JSON structure (fill with data from report, use null if not found):

{{
  "facility": {{
    "name": "facility name",
    "type": "CHC/PHC/etc",
    "block": "block name",
    "district": "district",
    "state": "state"
  }},
  "clinic_date": "YYYY-MM-DD or null",
  "beneficiaries": {{
    "expected": number or null,
    "attended": number or null,
    "attendance_rate": calculated percentage,
    "barriers": [
      {{
        "reason": "detailed reason",
        "count": number,
        "severity": "low/medium/high/critical",
        "root_cause": "analysis",
        "intervention": "what to do"
      }}
    ]
  }},
  "asha_performance": {{
    "name": "ASHA name",
    "home_visits": number or null,
    "issues": ["list of problems"],
    "rating": "poor/fair/good/excellent"
  }},
  "clinical_services": {{
    "staff_present": ["list of staff"],
    "examination_done": true/false,
    "counselling_topics": ["topics covered"],
    "counselling_gaps": ["topics missed"],
    "quality": "poor/fair/good/excellent"
  }},
  "laboratory": {{
    "tests_done": ["list of tests"],
    "sample_storage": "refrigerated/room_temp/etc",
    "cold_chain_maintained": true/false,
    "violations": ["any violations"],
    "turnaround_days": number or null
  }},
  "infrastructure_gaps": [
    {{
      "type": "space/equipment/staff/supplies",
      "description": "detailed description",
      "severity": "low/medium/high/critical",
      "impact": "how it affects service"
    }}
  ],
  "compliance": {{
    "due_list_prepared": true/false/null,
    "registers_updated": true/false/null,
    "protocol_deviations": ["list any deviations"],
    "score": 0-100
  }},
  "risks": [
    {{
      "risk": "description",
      "level": "low/medium/high/critical",
      "action_needed": "what to do",
      "timeline": "when to act"
    }}
  ],
  "recommendations": [
    {{
      "priority": 1-10,
      "action": "specific action",
      "responsible": "who should do it",
      "impact": "expected benefit"
    }}
  ],
  "executive_summary": "2-3 paragraph overall assessment",
  "key_findings": ["top 5 findings"],
  "critical_issues": ["urgent issues"],
  "overall_score": 0-100
}}

Return ONLY valid JSON, no other text."""

        print("🤖 Analyzing document with local AI model...")
        print("   (This may take 1-3 minutes depending on your model and hardware)")

        try:
            # Call local LM Studio API
            response = self.client.chat.completions.create(
                model=self.model_name,  # Use the configured model name
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert health data analyst. Extract structured data from reports and provide deep analysis. Always respond with valid JSON."
                    },
                    {
                        "role": "user",
                        "content": analysis_prompt
                    }
                ],
                temperature=0.3,
                max_tokens=8000,
            )

            response_text = response.choices[0].message.content
            print(f"\n📝 Raw AI response length: {len(response_text)} characters")

            # Try multiple JSON extraction strategies
            parsed_json = None

            # Strategy 1: Check for ```json blocks
            if "```json" in response_text:
                json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
                if json_match:
                    response_text = json_match.group(1)
            # Strategy 2: Check for ``` blocks
            elif "```" in response_text:
                json_match = re.search(r'```\s*(.*?)\s*```', response_text, re.DOTALL)
                if json_match:
                    response_text = json_match.group(1)

            # Strategy 3: Find JSON object by looking for { ... }
            if parsed_json is None:
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    response_text = json_match.group(0)

            # Clean up common issues
            response_text = response_text.strip()

            # Try to parse
            try:
                analysis = json.loads(response_text)
                print("✅ Successfully parsed JSON response")
                return analysis
            except json.JSONDecodeError as e:
                print(f"⚠️  First parse attempt failed: {e}")
                print(f"Response preview (first 500 chars):\n{response_text[:500]}...")
                print(f"\nResponse preview (last 500 chars):\n...{response_text[-500:]}")

                # Strategy 4: Try to fix common JSON errors
                # Remove trailing commas
                fixed_text = re.sub(r',(\s*[}\]])', r'\1', response_text)
                try:
                    analysis = json.loads(fixed_text)
                    print("✅ Successfully parsed after fixing trailing commas")
                    return analysis
                except json.JSONDecodeError:
                    pass

                # If all else fails, return minimal structure
                print("\n❌ Could not parse AI response as JSON after multiple attempts")
                return {
                    "error": "Failed to parse AI response",
                    "facility": {"name": "Unknown", "type": "Unknown"},
                    "executive_summary": "Analysis failed - could not parse AI response as JSON",
                    "overall_score": 0,
                    "raw_response_preview": response_text[:1000]
                }

        except Exception as e:
            print(f"❌ Unexpected error during analysis: {e}")
            import traceback
            traceback.print_exc()
            return {
                "error": f"Analysis error: {str(e)}",
                "facility": {"name": "Unknown", "type": "Unknown"},
                "executive_summary": "Analysis failed due to unexpected error",
                "overall_score": 0
            }

    def generate_report(self, analysis: Dict[str, Any], output_path: str):
        """Generate comprehensive markdown report."""

        report = f"""# 🏥 GOVERNMENT HEALTH DATA ANALYSIS REPORT
## Preconception & Maternal Health Clinic (PPC) Assessment

**Report Generated**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
**Analysis System**: Local AI-Powered Analysis (LM Studio)

---

# 📊 EXECUTIVE SUMMARY

{analysis.get('executive_summary', 'No summary available')}

**Overall Score**: {analysis.get('overall_score', 'N/A')}/100

## 🎯 KEY FINDINGS

"""

        for i, finding in enumerate(analysis.get('key_findings', []), 1):
            report += f"{i}. **{finding}**\n"

        report += "\n"

        if analysis.get('critical_issues'):
            report += "## 🚨 CRITICAL ISSUES\n\n"
            for issue in analysis['critical_issues']:
                report += f"- ⚠️ **{issue}**\n"
            report += "\n"

        report += "---\n\n"

        # Facility
        fac = analysis.get('facility', {})
        report += f"""# 🏢 FACILITY INFORMATION

**Name**: {fac.get('name', 'Unknown')}
**Type**: {fac.get('type', 'Unknown')}
**Location**: {fac.get('block', 'Unknown')}, {fac.get('district', 'Unknown')}, {fac.get('state', 'Unknown')}

**Clinic Date**: {analysis.get('clinic_date', 'Not specified')}

---

"""

        # Beneficiaries
        ben = analysis.get('beneficiaries', {})
        if ben:
            report += f"""# 👥 BENEFICIARY ATTENDANCE

**Expected**: {ben.get('expected', 'Unknown')}
**Attended**: {ben.get('attended', 'Unknown')}
**Attendance Rate**: {ben.get('attendance_rate', 'Unknown')}

"""

            if ben.get('barriers'):
                report += "## 🚧 ATTENDANCE BARRIERS\n\n"
                for barrier in ben['barriers']:
                    report += f"""### {barrier.get('reason', 'Unknown')}
- **Count**: {barrier.get('count', 'Unknown')} beneficiaries
- **Severity**: {barrier.get('severity', 'Unknown').upper()}
- **Root Cause**: {barrier.get('root_cause', 'Not analyzed')}
- **Intervention**: {barrier.get('intervention', 'None specified')}

"""

        # ASHA Performance
        asha = analysis.get('asha_performance', {})
        if asha:
            report += f"""## 👩‍⚕️ ASHA PERFORMANCE

**Name**: {asha.get('name', 'Unknown')}
**Home Visits**: {asha.get('home_visits', 'Unknown')}
**Rating**: {asha.get('rating', 'Unknown').upper()}

### Issues Identified
"""
            for issue in asha.get('issues', []):
                report += f"- ❌ {issue}\n"
            report += "\n---\n\n"

        # Clinical Services
        clin = analysis.get('clinical_services', {})
        if clin:
            report += f"""# 💉 CLINICAL SERVICES

**Quality Rating**: {clin.get('quality', 'Unknown').upper()}

**Staff Present**: {', '.join(clin.get('staff_present', []))}

**Examination Done**: {clin.get('examination_done', 'Unknown')}

## Counselling
**Topics Covered**: {', '.join(clin.get('counselling_topics', []))}

**Gaps**: {', '.join(clin.get('counselling_gaps', []))}

---

"""

        # Laboratory
        lab = analysis.get('laboratory', {})
        if lab:
            report += f"""# 🧪 LABORATORY SERVICES

**Tests Done**: {', '.join(lab.get('tests_done', []))}

**Sample Storage**: {lab.get('sample_storage', 'Unknown').title()}

**Cold Chain Maintained**: {lab.get('cold_chain_maintained', 'Unknown')}

**Turnaround Time**: {lab.get('turnaround_days', 'Unknown')} days

"""
            if lab.get('violations'):
                report += "### ⚠️ VIOLATIONS\n"
                for violation in lab['violations']:
                    report += f"- ⛔ {violation}\n"
                report += "\n"

            report += "---\n\n"

        # Infrastructure Gaps
        if analysis.get('infrastructure_gaps'):
            report += "# 🏗️ INFRASTRUCTURE GAPS\n\n"
            for gap in analysis['infrastructure_gaps']:
                report += f"""### {gap.get('type', 'Unknown').upper()} - {gap.get('severity', 'Unknown').upper()}

**Description**: {gap.get('description', 'No description')}

**Impact**: {gap.get('impact', 'Not assessed')}

---

"""

        # Compliance
        comp = analysis.get('compliance', {})
        if comp:
            report += f"""# 📋 COMPLIANCE

**Score**: {comp.get('score', 'N/A')}/100

- **Due List Prepared**: {comp.get('due_list_prepared', 'Unknown')}
- **Registers Updated**: {comp.get('registers_updated', 'Unknown')}

"""
            if comp.get('protocol_deviations'):
                report += "### Protocol Deviations\n"
                for dev in comp['protocol_deviations']:
                    report += f"- {dev}\n"
                report += "\n"

            report += "---\n\n"

        # Risks
        if analysis.get('risks'):
            report += "# ⚠️ RISK ASSESSMENT\n\n"
            for risk in analysis['risks']:
                report += f"""### {risk.get('risk', 'Unknown')} - {risk.get('level', 'Unknown').upper()}

**Action Needed**: {risk.get('action_needed', 'Not specified')}

**Timeline**: {risk.get('timeline', 'Not specified')}

---

"""

        # Recommendations
        if analysis.get('recommendations'):
            report += "# 💡 RECOMMENDATIONS\n\n"
            for rec in sorted(analysis['recommendations'], key=lambda x: x.get('priority', 99)):
                report += f"""### Priority {rec.get('priority', '?')}: {rec.get('action', 'Unknown')}

**Responsible**: {rec.get('responsible', 'Not assigned')}

**Expected Impact**: {rec.get('impact', 'Not specified')}

---

"""

        report += f"""
---

# CONCLUSION

This analysis was generated using local AI models running on LM Studio, demonstrating:
- Extraction of structured data from unpredictable input
- Deep analysis of health service delivery
- Risk identification and prioritized recommendations
- Complete offline operation with no cloud dependency

**Analysis completed locally - no data sent to external servers.**

---

**Report End**
**Generated by**: Government Health Data Automation Platform (Local)
**Technology**: LM Studio Local AI
**Date**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
"""

        # Write report
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report)

        print(f"\n✅ Report generated: {output_path}")

        # Save JSON
        json_path = output_path.replace('.md', '.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, indent=2, ensure_ascii=False)

        print(f"✅ JSON saved: {json_path}")


def main():
    """Main entry point."""

    print("""
╔══════════════════════════════════════════════════════════════════════════════╗
║                                                                              ║
║     🏥 GOVERNMENT HEALTH DATA AUTOMATION - LOCAL AI VERSION 🏥               ║
║                                                                              ║
║              Powered by LM Studio (100% Offline Analysis)                    ║
║                                                                              ║
╚══════════════════════════════════════════════════════════════════════════════╝
""")

    if len(sys.argv) != 2:
        print("Usage: python poc_analyzer_local.py <path_to_report.docx>")
        print("\nExample: python poc_analyzer_local.py SAMPLE_REPORT.txt")
        print("\nMake sure LM Studio is running with a loaded model!")
        sys.exit(1)

    input_file = sys.argv[1]

    if not Path(input_file).exists():
        print(f"❌ ERROR: File not found: {input_file}")
        sys.exit(1)

    # Check for custom LM Studio URL
    lm_studio_url = os.environ.get("LM_STUDIO_URL", "http://192.168.56.1:1234/v1")

    try:
        print(f"🔌 Connecting to LM Studio at: {lm_studio_url}")
        analyzer = LocalGovernmentHealthAnalyzer(base_url=lm_studio_url)

        print(f"📄 Extracting text from: {input_file}")
        raw_text = analyzer.extract_text(input_file)

        print(f"✅ Extracted {len(raw_text)} characters")

        print("\n" + "="*80)
        analysis = analyzer.analyze_document(raw_text)

        output_path = Path(input_file).stem + "_LOCAL_ANALYSIS.md"
        print(f"\n📊 Generating report...")
        analyzer.generate_report(analysis, output_path)

        print("\n" + "="*80)
        print("\n🎉 SUCCESS! Local AI analysis complete.")
        print(f"\n📄 Reports generated:")
        print(f"   - Markdown: {output_path}")
        print(f"   - JSON: {output_path.replace('.md', '.json')}")

        print("\n✨ Advantages of local processing:")
        print("   ✅ 100% offline - no data leaves your machine")
        print("   ✅ No API costs - unlimited analysis")
        print("   ✅ Data privacy - government-safe")
        print("   ✅ No internet dependency")

        print("\n💡 For best results, use models like:")
        print("   - Qwen 2.5 (14B or 32B)")
        print("   - Llama 3.1 (8B or 70B)")
        print("   - Mistral (7B or larger)")
        print("\n")

    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        print("\n💡 Troubleshooting:")
        print("   1. Is LM Studio running?")
        print("   2. Is a model loaded in LM Studio?")
        print("   3. Is the local server started?")
        print(f"   4. Can you access: {lm_studio_url}")
        print("   5. Check the model name - look in LM Studio for the exact name")
        print("      Set it with: set LM_STUDIO_MODEL=your-model-name")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
