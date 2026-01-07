#!/usr/bin/env python3
"""
Government Health Data Automation - Enhanced Local Analyzer with Beautiful PDF
Generates professional, stakeholder-ready PDF reports with excellent UI/UX

Usage:
    python poc_analyzer_local_enhanced.py path/to/report.docx

Requirements:
    pip install openai python-docx PyPDF2 python-dateutil reportlab pillow
"""

import os
import re
import sys
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional

# Document processing imports
try:
    import docx
    from PyPDF2 import PdfReader
except ImportError:
    print("ERROR: Required packages not installed.")
    print("Install with: pip install openai python-docx PyPDF2 python-dateutil reportlab pillow")
    sys.exit(1)

# PDF generation imports
try:
    from reportlab.lib import colors
    from reportlab.lib.units import inch, cm
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        PageBreak, Image, KeepTogether, Frame, PageTemplate,
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )
    from reportlab.pdfgen import canvas
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
except ImportError:
    print("ERROR: ReportLab not installed.")
    print("Install with: pip install reportlab pillow")
    sys.exit(1)

# OpenAI-compatible client for LM Studio
try:
    from openai import OpenAI
except ImportError:
    print("ERROR: OpenAI package not installed.")
    print("Install with: pip install openai")
    sys.exit(1)


class BeautifulPDFGenerator:
    """Generate professional, award-winning PDF reports."""

    def __init__(self, output_path: str):
        """Initialize PDF generator."""
        self.output_path = output_path
        self.story = []
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Create custom paragraph styles for beautiful formatting."""

        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a237e'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))

        # Subtitle style
        self.styles.add(ParagraphStyle(
            name='CustomSubtitle',
            parent=self.styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#3949ab'),
            spaceAfter=20,
            spaceBefore=10,
            alignment=TA_CENTER,
            fontName='Helvetica'
        ))

        # Section header
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=18,
            textColor=colors.HexColor('#1565c0'),
            spaceAfter=12,
            spaceBefore=20,
            fontName='Helvetica-Bold',
            borderWidth=0,
            borderColor=colors.HexColor('#1565c0'),
            borderPadding=5,
            leftIndent=0
        ))

        # Subsection header
        self.styles.add(ParagraphStyle(
            name='SubsectionHeader',
            parent=self.styles['Heading3'],
            fontSize=14,
            textColor=colors.HexColor('#0277bd'),
            spaceAfter=10,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        ))

        # Body text
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['BodyText'],
            fontSize=11,
            textColor=colors.HexColor('#212121'),
            alignment=TA_JUSTIFY,
            spaceAfter=8,
            leading=16
        ))

        # Highlight box
        self.styles.add(ParagraphStyle(
            name='HighlightBox',
            parent=self.styles['BodyText'],
            fontSize=11,
            textColor=colors.HexColor('#c62828'),
            backColor=colors.HexColor('#ffebee'),
            borderWidth=1,
            borderColor=colors.HexColor('#c62828'),
            borderPadding=10,
            spaceAfter=15,
            leading=16
        ))

        # Success box
        self.styles.add(ParagraphStyle(
            name='SuccessBox',
            parent=self.styles['BodyText'],
            fontSize=11,
            textColor=colors.HexColor('#2e7d32'),
            backColor=colors.HexColor('#e8f5e9'),
            borderWidth=1,
            borderColor=colors.HexColor('#2e7d32'),
            borderPadding=10,
            spaceAfter=15,
            leading=16
        ))

        # Quote style
        self.styles.add(ParagraphStyle(
            name='Quote',
            parent=self.styles['BodyText'],
            fontSize=10,
            textColor=colors.HexColor('#424242'),
            leftIndent=20,
            rightIndent=20,
            fontName='Helvetica-Oblique',
            spaceAfter=10
        ))

        # Footer style
        self.styles.add(ParagraphStyle(
            name='Footer',
            parent=self.styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#757575'),
            alignment=TA_CENTER
        ))

    def add_cover_page(self, analysis: Dict[str, Any], facility_name: str, report_date: str):
        """Add professional cover page."""

        # Title
        self.story.append(Spacer(1, 2*inch))

        title = Paragraph(
            "🏥 GOVERNMENT HEALTH CLINIC<br/>ANALYSIS REPORT",
            self.styles['CustomTitle']
        )
        self.story.append(title)
        self.story.append(Spacer(1, 0.3*inch))

        # Facility name
        facility_para = Paragraph(
            f"<b>{facility_name}</b>",
            self.styles['CustomSubtitle']
        )
        self.story.append(facility_para)
        self.story.append(Spacer(1, 0.2*inch))

        # Report date
        date_para = Paragraph(
            f"Clinic Date: {report_date}",
            self.styles['CustomSubtitle']
        )
        self.story.append(date_para)
        self.story.append(Spacer(1, 1*inch))

        # Overall score box
        score = analysis.get('overall_score', 'N/A')
        score_color = self._get_score_color(score)

        score_data = [[
            Paragraph(f"<b>OVERALL PERFORMANCE SCORE</b>", self.styles['CustomBody']),
            Paragraph(f"<b><font size=24 color='{score_color}'>{score}/100</font></b>",
                     self.styles['CustomBody'])
        ]]

        score_table = Table(score_data, colWidths=[3*inch, 2*inch])
        score_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#1565c0')),
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#e3f2fd')),
            ('PADDING', (0, 0), (-1, -1), 15),
        ]))

        self.story.append(score_table)
        self.story.append(Spacer(1, 0.5*inch))

        # Analysis metadata
        metadata_text = f"""
        <b>Report Generated:</b> {datetime.now().strftime("%d %B %Y, %I:%M %p")}<br/>
        <b>Analysis System:</b> AI-Powered Health Data Automation<br/>
        <b>Technology:</b> Local AI Processing (100% Offline)<br/>
        <b>Classification:</b> For Official Use
        """

        metadata_para = Paragraph(metadata_text, self.styles['CustomBody'])
        self.story.append(metadata_para)

        self.story.append(PageBreak())

    def _get_score_color(self, score):
        """Get color based on score."""
        if isinstance(score, str) or score is None:
            return '#757575'

        if score >= 80:
            return '#2e7d32'  # Green
        elif score >= 60:
            return '#f57c00'  # Orange
        else:
            return '#c62828'  # Red

    def _get_severity_color(self, severity: str):
        """Get color based on severity."""
        severity_colors = {
            'critical': '#c62828',
            'high': '#e64a19',
            'medium': '#f57c00',
            'low': '#558b2f'
        }
        return severity_colors.get(severity.lower(), '#757575')

    def add_executive_summary(self, analysis: Dict[str, Any]):
        """Add executive summary section."""

        self.story.append(Paragraph("📊 EXECUTIVE SUMMARY", self.styles['SectionHeader']))
        self.story.append(Spacer(1, 0.2*inch))

        # One-sentence summary
        summary_text = analysis.get('executive_summary', 'Analysis summary not available.')
        summary_para = Paragraph(summary_text, self.styles['CustomBody'])
        self.story.append(summary_para)
        self.story.append(Spacer(1, 0.15*inch))

        # Key findings
        if analysis.get('key_findings'):
            self.story.append(Paragraph("🎯 Key Findings", self.styles['SubsectionHeader']))

            for i, finding in enumerate(analysis['key_findings'][:7], 1):
                finding_text = f"<b>{i}.</b> {finding}"
                self.story.append(Paragraph(finding_text, self.styles['CustomBody']))
                self.story.append(Spacer(1, 0.05*inch))

        # Critical issues (highlighted)
        if analysis.get('critical_issues'):
            self.story.append(Spacer(1, 0.15*inch))
            self.story.append(Paragraph("🚨 CRITICAL ISSUES REQUIRING IMMEDIATE ATTENTION",
                                       self.styles['SubsectionHeader']))

            for issue in analysis['critical_issues']:
                issue_para = Paragraph(f"⚠️ <b>{issue}</b>", self.styles['HighlightBox'])
                self.story.append(issue_para)

        self.story.append(PageBreak())

    def add_facility_info(self, analysis: Dict[str, Any]):
        """Add facility information section."""

        self.story.append(Paragraph("🏢 FACILITY INFORMATION", self.styles['SectionHeader']))
        self.story.append(Spacer(1, 0.15*inch))

        facility = analysis.get('facility', {})
        clinic_date = analysis.get('clinic_date', 'Not specified')

        # Facility details table
        data = [
            ['Facility Name', facility.get('name', 'Unknown')],
            ['Facility Type', facility.get('type', 'Unknown')],
            ['Block', facility.get('block', 'Unknown')],
            ['District', facility.get('district', 'Unknown')],
            ['State', facility.get('state', 'Unknown')],
            ['Clinic Date', clinic_date],
        ]

        table = Table(data, colWidths=[2*inch, 4*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e3f2fd')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#212121')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#90caf9')),
            ('PADDING', (0, 0), (-1, -1), 10),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))

        self.story.append(table)
        self.story.append(Spacer(1, 0.2*inch))

    def add_beneficiary_analysis(self, analysis: Dict[str, Any]):
        """Add beneficiary attendance analysis."""

        self.story.append(PageBreak())
        self.story.append(Paragraph("👥 BENEFICIARY ATTENDANCE ANALYSIS",
                                   self.styles['SectionHeader']))
        self.story.append(Spacer(1, 0.15*inch))

        ben = analysis.get('beneficiaries', {})

        # Attendance metrics
        expected = ben.get('expected', 'Unknown')
        attended = ben.get('attended', 'Unknown')
        rate = ben.get('attendance_rate', 'Unknown')

        if isinstance(rate, (int, float)):
            rate_display = f"{rate:.1%}"
        else:
            rate_display = str(rate)

        metrics_data = [
            ['Expected Beneficiaries', str(expected)],
            ['Actually Attended', str(attended)],
            ['Attendance Rate', rate_display],
        ]

        metrics_table = Table(metrics_data, colWidths=[2.5*inch, 1.5*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e8f5e9')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#212121')),
            ('ALIGN', (1, 0), (1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#66bb6a')),
            ('PADDING', (0, 0), (-1, -1), 12),
        ]))

        self.story.append(metrics_table)
        self.story.append(Spacer(1, 0.2*inch))

        # Attendance barriers
        if ben.get('barriers'):
            self.story.append(Paragraph("🚧 Attendance Barriers (Deep Analysis)",
                                       self.styles['SubsectionHeader']))
            self.story.append(Spacer(1, 0.1*inch))

            for i, barrier in enumerate(ben['barriers'][:5], 1):
                severity = barrier.get('severity', 'unknown')
                severity_color = self._get_severity_color(severity)

                barrier_header = f"<b>Barrier #{i}: {barrier.get('reason', 'Unknown')}</b>"
                self.story.append(Paragraph(barrier_header, self.styles['SubsectionHeader']))

                barrier_details = f"""
                <b>Affected Beneficiaries:</b> {barrier.get('count', 'Unknown')}<br/>
                <b>Severity:</b> <font color='{severity_color}'><b>{severity.upper()}</b></font><br/>
                <b>Root Cause:</b> {barrier.get('root_cause', 'Not analyzed')}<br/>
                <b>Recommended Intervention:</b> {barrier.get('intervention', 'None specified')}
                """

                self.story.append(Paragraph(barrier_details, self.styles['CustomBody']))
                self.story.append(Spacer(1, 0.15*inch))

    def add_clinical_services(self, analysis: Dict[str, Any]):
        """Add clinical services section."""

        self.story.append(PageBreak())
        self.story.append(Paragraph("💉 CLINICAL SERVICES ASSESSMENT",
                                   self.styles['SectionHeader']))
        self.story.append(Spacer(1, 0.15*inch))

        clin = analysis.get('clinical_services', {})

        # Quality rating
        quality = clin.get('quality', 'Unknown')
        quality_color = self._get_score_color(
            {'poor': 25, 'fair': 50, 'good': 75, 'excellent': 95}.get(quality.lower(), 0)
        )

        quality_text = f"<b>Overall Service Quality:</b> <font color='{quality_color}'><b>{quality.upper()}</b></font>"
        self.story.append(Paragraph(quality_text, self.styles['CustomBody']))
        self.story.append(Spacer(1, 0.15*inch))

        # Staff present
        if clin.get('staff_present'):
            self.story.append(Paragraph("Staff Present", self.styles['SubsectionHeader']))
            staff_text = ", ".join(clin['staff_present'])
            self.story.append(Paragraph(staff_text, self.styles['CustomBody']))
            self.story.append(Spacer(1, 0.1*inch))

        # Counselling
        if clin.get('counselling_topics'):
            self.story.append(Paragraph("Counselling Provided", self.styles['SubsectionHeader']))
            topics_text = "✅ " + "<br/>✅ ".join(clin['counselling_topics'])
            self.story.append(Paragraph(topics_text, self.styles['CustomBody']))
            self.story.append(Spacer(1, 0.1*inch))

        # Gaps
        if clin.get('counselling_gaps'):
            gaps_text = "❌ <b>Gaps:</b> " + ", ".join(clin['counselling_gaps'])
            self.story.append(Paragraph(gaps_text, self.styles['HighlightBox']))

    def add_laboratory_services(self, analysis: Dict[str, Any]):
        """Add laboratory services section."""

        lab = analysis.get('laboratory', {})
        if not lab:
            return

        self.story.append(PageBreak())
        self.story.append(Paragraph("🧪 LABORATORY SERVICES", self.styles['SectionHeader']))
        self.story.append(Spacer(1, 0.15*inch))

        # Tests conducted
        if lab.get('tests_done'):
            self.story.append(Paragraph("Tests Conducted", self.styles['SubsectionHeader']))
            tests_text = "• " + "<br/>• ".join(lab['tests_done'])
            self.story.append(Paragraph(tests_text, self.styles['CustomBody']))
            self.story.append(Spacer(1, 0.15*inch))

        # Sample handling
        storage = lab.get('sample_storage', 'Unknown')
        cold_chain = lab.get('cold_chain_maintained', None)

        storage_text = f"<b>Sample Storage:</b> {storage.title()}<br/>"
        storage_text += f"<b>Cold Chain Maintained:</b> "

        if cold_chain:
            storage_text += "<font color='#2e7d32'>✅ YES</font>"
        elif cold_chain is False:
            storage_text += "<font color='#c62828'>⛔ NO</font>"
        else:
            storage_text += "Unknown"

        self.story.append(Paragraph(storage_text, self.styles['CustomBody']))
        self.story.append(Spacer(1, 0.1*inch))

        # Violations (critical highlighting)
        if lab.get('violations'):
            self.story.append(Paragraph("⚠️ CRITICAL VIOLATIONS DETECTED",
                                       self.styles['SubsectionHeader']))
            for violation in lab['violations']:
                violation_para = Paragraph(f"⛔ {violation}", self.styles['HighlightBox'])
                self.story.append(violation_para)

    def add_risk_assessment(self, analysis: Dict[str, Any]):
        """Add risk assessment section."""

        risks = analysis.get('risks', [])
        if not risks:
            return

        self.story.append(PageBreak())
        self.story.append(Paragraph("⚠️ RISK ASSESSMENT", self.styles['SectionHeader']))
        self.story.append(Spacer(1, 0.15*inch))

        for risk in risks[:5]:
            level = risk.get('level', 'unknown')
            level_color = self._get_severity_color(level)

            risk_header = f"<b>{risk.get('risk', 'Unknown Risk')}</b>"
            self.story.append(Paragraph(risk_header, self.styles['SubsectionHeader']))

            risk_details = f"""
            <b>Risk Level:</b> <font color='{level_color}'><b>{level.upper()}</b></font><br/>
            <b>Action Needed:</b> {risk.get('action_needed', 'Not specified')}<br/>
            <b>Timeline:</b> {risk.get('timeline', 'Not specified')}
            """

            self.story.append(Paragraph(risk_details, self.styles['CustomBody']))
            self.story.append(Spacer(1, 0.15*inch))

    def add_recommendations(self, analysis: Dict[str, Any]):
        """Add actionable recommendations section."""

        recs = analysis.get('recommendations', [])
        if not recs:
            return

        self.story.append(PageBreak())
        self.story.append(Paragraph("💡 ACTIONABLE RECOMMENDATIONS",
                                   self.styles['SectionHeader']))
        self.story.append(Spacer(1, 0.15*inch))

        # Sort by priority
        sorted_recs = sorted(recs, key=lambda x: x.get('priority', 99))

        for rec in sorted_recs[:10]:
            priority = rec.get('priority', '?')

            # Priority badge color
            if isinstance(priority, int):
                if priority <= 3:
                    priority_color = '#c62828'
                elif priority <= 6:
                    priority_color = '#f57c00'
                else:
                    priority_color = '#558b2f'
            else:
                priority_color = '#757575'

            rec_header = f"<b>Priority {priority}:</b> {rec.get('action', 'Unknown action')}"
            self.story.append(Paragraph(rec_header, self.styles['SubsectionHeader']))

            rec_details = f"""
            <b>Responsible Party:</b> {rec.get('responsible', 'Not assigned')}<br/>
            <b>Expected Impact:</b> {rec.get('impact', 'Not specified')}
            """

            self.story.append(Paragraph(rec_details, self.styles['CustomBody']))
            self.story.append(Spacer(1, 0.12*inch))

    def add_footer(self):
        """Add closing footer."""

        self.story.append(PageBreak())
        self.story.append(Spacer(1, 2*inch))

        footer_text = f"""
        <b>— END OF REPORT —</b><br/><br/>
        This comprehensive analysis demonstrates the capability of AI-powered automation
        to extract deep, actionable insights from field reports, regardless of data quality
        or format.<br/><br/>
        <b>Generated by:</b> Government Health Data Automation Platform<br/>
        <b>Technology:</b> Local AI Processing (100% Offline)<br/>
        <b>Date:</b> {datetime.now().strftime("%d %B %Y, %I:%M %p")}<br/><br/>
        <i>For Official Use Only - Confidential</i>
        """

        self.story.append(Paragraph(footer_text, self.styles['Footer']))

    def generate(self, analysis: Dict[str, Any], facility_name: str, report_date: str):
        """Generate complete PDF report."""

        # Create PDF
        doc = SimpleDocTemplate(
            self.output_path,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )

        # Add all sections
        self.add_cover_page(analysis, facility_name, report_date)
        self.add_executive_summary(analysis)
        self.add_facility_info(analysis)
        self.add_beneficiary_analysis(analysis)
        self.add_clinical_services(analysis)
        self.add_laboratory_services(analysis)
        self.add_risk_assessment(analysis)
        self.add_recommendations(analysis)
        self.add_footer()

        # Build PDF
        doc.build(self.story)


class LocalGovernmentHealthAnalyzer:
    """Enhanced local analyzer with beautiful PDF generation and organized output."""

    def __init__(self, base_url: str = "http://192.168.56.1:1234/v1", model_name: str = None):
        """Initialize analyzer."""
        self.base_url = base_url
        self.model_name = model_name or os.environ.get("LM_STUDIO_MODEL", "openai/gpt-oss-20b")
        self.client = OpenAI(base_url=self.base_url, api_key="lm-studio")
        print(f"✅ Connected to local LM Studio at: {self.base_url}")
        print(f"🤖 Using model: {self.model_name}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        doc = docx.Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                text_parts.append(" | ".join(row_data))

        return "\n".join(text_parts)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)

        return "\n".join(text_parts)

    def extract_text(self, file_path: str) -> str:
        """Extract text from document."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension == ".docx":
            return self.extract_text_from_docx(str(file_path))
        elif extension == ".pdf":
            return self.extract_text_from_pdf(str(file_path))
        else:
            return file_path.read_text(encoding="utf-8", errors="ignore")

    def analyze_document(self, raw_text: str) -> Dict[str, Any]:
        """AI-powered analysis using local LM Studio."""

        analysis_prompt = f"""You are a government health data analyst for Maternal Health Clinics in India.

Analyze this field report and extract ALL information into structured JSON.

REPORT:
{raw_text}

---

Extract this JSON structure (fill with data from report, use null if not found):

{{
  "facility": {{
    "name": "facility name",
    "type": "CHC/PHC/etc",
    "block": "block name",
    "district": "district",
    "state": "state"
  }},
  "clinic_date": "YYYY-MM-DD or null",
  "beneficiaries": {{
    "expected": number or null,
    "attended": number or null,
    "attendance_rate": calculated percentage (0-1),
    "barriers": [
      {{
        "reason": "detailed reason from report",
        "count": number,
        "severity": "low/medium/high/critical",
        "root_cause": "deep analysis of why this happened",
        "intervention": "specific actionable recommendation"
      }}
    ]
  }},
  "asha_performance": {{
    "name": "ASHA name",
    "home_visits": number or null,
    "issues": ["list of problems"],
    "rating": "poor/fair/good/excellent"
  }},
  "clinical_services": {{
    "staff_present": ["list of staff with designations"],
    "examination_done": true/false,
    "counselling_topics": ["topics covered like nutrition, exercise, etc"],
    "counselling_gaps": ["topics that should have been covered but weren't"],
    "quality": "poor/fair/good/excellent"
  }},
  "laboratory": {{
    "tests_done": ["list of tests conducted"],
    "sample_storage": "refrigerated/room_temp/ice_box/etc",
    "cold_chain_maintained": true/false/null,
    "violations": ["any critical violations like improper storage"],
    "turnaround_days": number or null
  }},
  "infrastructure_gaps": [
    {{
      "type": "space/equipment/staff/supplies",
      "description": "detailed description",
      "severity": "low/medium/high/critical",
      "impact": "how it affects service delivery"
    }}
  ],
  "compliance": {{
    "due_list_prepared": true/false/null,
    "registers_updated": true/false/null,
    "protocol_deviations": ["list any deviations from standard protocols"],
    "score": estimated score 0-100
  }},
  "risks": [
    {{
      "risk": "description of risk",
      "level": "low/medium/high/critical",
      "action_needed": "what must be done",
      "timeline": "when to act (immediate/short-term/long-term)"
    }}
  ],
  "recommendations": [
    {{
      "priority": 1-10 (1=highest),
      "action": "specific actionable recommendation",
      "responsible": "who should do it (ASHA/Medical Officer/Block Coordinator/etc)",
      "impact": "expected benefit if implemented"
    }}
  ],
  "executive_summary": "comprehensive 2-3 paragraph assessment covering main findings, critical issues, and overall status",
  "key_findings": ["top 5-7 most important findings as bullet points"],
  "critical_issues": ["issues requiring immediate attention"],
  "overall_score": estimated overall performance score 0-100
}}

Important:
- Extract EVERYTHING you can from the text
- Provide deep root cause analysis for barriers
- Make recommendations specific and actionable
- Estimate scores based on performance indicators
- Be thorough in the executive summary

Return ONLY valid JSON, no other text."""

        print("🤖 Analyzing document with local AI model...")
        print("   (This may take 1-3 minutes depending on your model and hardware)")
        print(f"   📏 Document length: {len(raw_text)} characters")

        # Warn if document is very large
        if len(raw_text) > 15000:
            print(f"   ⚠️  WARNING: Large document may exceed model's context window")
            print(f"   💡 Consider using a model with larger context (Qwen 2.5, Llama 3.1)")

        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert health data analyst. Extract structured data from reports and provide deep analysis. Always respond with valid JSON."
                    },
                    {
                        "role": "user",
                        "content": analysis_prompt
                    }
                ],
                temperature=0.3,
                max_tokens=8000,
                timeout=300.0  # 5 minute timeout
            )

            response_text = response.choices[0].message.content
            print(f"\n📝 Raw AI response length: {len(response_text)} characters")

            # Check for empty response
            if not response_text or len(response_text) == 0:
                print("\n❌ CRITICAL: Model returned empty response!")
                print("💡 Possible causes:")
                print("   1. Document too long for model's context window")
                print("   2. Content filtering/safety filter triggered")
                print("   3. Model timeout or crash")
                print("   4. LM Studio configuration issue")
                print("\n💡 Solutions:")
                print("   1. Try a different model (Qwen 2.5 recommended)")
                print("   2. Use a smaller document for testing")
                print("   3. Check LM Studio console for errors")
                raise ValueError("Model returned empty response")

            # Try multiple JSON extraction strategies
            parsed_json = None

            # Strategy 1: Check for ```json blocks
            if "```json" in response_text:
                json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
                if json_match:
                    response_text = json_match.group(1)
            # Strategy 2: Check for ``` blocks
            elif "```" in response_text:
                json_match = re.search(r'```\s*(.*?)\s*```', response_text, re.DOTALL)
                if json_match:
                    response_text = json_match.group(1)

            # Strategy 3: Find JSON object by looking for { ... }
            if parsed_json is None:
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    response_text = json_match.group(0)

            # Clean up common issues
            response_text = response_text.strip()

            # Try to parse
            try:
                analysis = json.loads(response_text)
                print("✅ Successfully parsed JSON response")
                return analysis
            except json.JSONDecodeError as e:
                print(f"⚠️  First parse attempt failed: {e}")
                print(f"Response preview (first 500 chars):\n{response_text[:500]}...")
                print(f"\nResponse preview (last 500 chars):\n...{response_text[-500:]}")

                # Strategy 4: Try to fix common JSON errors
                # Remove trailing commas
                fixed_text = re.sub(r',(\s*[}\]])', r'\1', response_text)
                try:
                    analysis = json.loads(fixed_text)
                    print("✅ Successfully parsed after fixing trailing commas")
                    return analysis
                except json.JSONDecodeError:
                    pass

                # If all else fails, return minimal structure
                print("\n❌ Could not parse AI response as JSON after multiple attempts")
                return {
                    "error": "Failed to parse AI response",
                    "facility": {"name": "Unknown", "type": "Unknown"},
                    "executive_summary": "Analysis failed - could not parse AI response as JSON",
                    "overall_score": 0,
                    "raw_response_preview": response_text[:1000]
                }

        except Exception as e:
            print(f"❌ Unexpected error during analysis: {e}")
            import traceback
            traceback.print_exc()
            return {
                "error": f"Analysis error: {str(e)}",
                "facility": {"name": "Unknown", "type": "Unknown"},
                "executive_summary": "Analysis failed due to unexpected error",
                "overall_score": 0
            }


def main():
    """Main entry point."""

    print("""
╔══════════════════════════════════════════════════════════════════════════════╗
║                                                                              ║
║        🏥 GOVERNMENT HEALTH DATA AUTOMATION - ENHANCED LOCAL 🏥              ║
║                                                                              ║
║           Beautiful PDF Reports + Organized Output Folders                   ║
║                                                                              ║
╚══════════════════════════════════════════════════════════════════════════════╝
""")

    if len(sys.argv) != 2:
        print("Usage: python poc_analyzer_local_enhanced.py <path_to_report.docx>")
        print("\nExample: python poc_analyzer_local_enhanced.py SAMPLE_REPORT.txt")
        sys.exit(1)

    input_file = sys.argv[1]

    if not Path(input_file).exists():
        print(f"❌ ERROR: File not found: {input_file}")
        sys.exit(1)

    lm_studio_url = os.environ.get("LM_STUDIO_URL", "http://192.168.56.1:1234/v1")

    try:
        # Create output folder with timestamp
        input_path = Path(input_file)
        file_stem = input_path.stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_folder = Path(f"analysis_reports/{file_stem}_{timestamp}")
        output_folder.mkdir(parents=True, exist_ok=True)

        print(f"📁 Output folder: {output_folder}")
        print(f"🔌 Connecting to LM Studio at: {lm_studio_url}")

        # Initialize analyzer
        analyzer = LocalGovernmentHealthAnalyzer(base_url=lm_studio_url)

        # Extract text
        print(f"📄 Extracting text from: {input_file}")
        raw_text = analyzer.extract_text(input_file)
        print(f"✅ Extracted {len(raw_text)} characters")

        # Save raw text
        raw_text_path = output_folder / "01_raw_text.txt"
        raw_text_path.write_text(raw_text, encoding='utf-8')
        print(f"✅ Saved raw text: {raw_text_path.name}")

        # Analyze
        print("\n" + "="*80)
        analysis = analyzer.analyze_document(raw_text)

        # Save JSON
        json_path = output_folder / "02_analysis_data.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, indent=2, ensure_ascii=False)
        print(f"✅ Saved JSON data: {json_path.name}")

        # Generate markdown report
        md_path = output_folder / "03_detailed_report.md"

        facility_name = analysis.get('facility', {}).get('name', 'Unknown Facility')
        clinic_date = analysis.get('clinic_date', 'Not specified')

        # Simple markdown generation (keeping existing logic)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(f"# Health Clinic Analysis Report\n\n")
            f.write(f"**Facility**: {facility_name}\n")
            f.write(f"**Date**: {clinic_date}\n")
            f.write(f"**Overall Score**: {analysis.get('overall_score', 'N/A')}/100\n\n")
            f.write(f"## Executive Summary\n\n{analysis.get('executive_summary', 'N/A')}\n\n")

            if analysis.get('key_findings'):
                f.write("## Key Findings\n\n")
                for i, finding in enumerate(analysis['key_findings'], 1):
                    f.write(f"{i}. {finding}\n")
                f.write("\n")

        print(f"✅ Saved markdown report: {md_path.name}")

        # Generate beautiful PDF
        print("\n📊 Generating beautiful PDF report...")
        pdf_path = output_folder / "04_STAKEHOLDER_REPORT.pdf"

        pdf_gen = BeautifulPDFGenerator(str(pdf_path))
        pdf_gen.generate(analysis, facility_name, clinic_date)

        print(f"✅ Saved beautiful PDF: {pdf_path.name}")

        # Print summary
        print("\n" + "="*80)
        print("\n🎉 SUCCESS! Analysis complete.")
        print(f"\n📁 All files saved in: {output_folder}")
        print(f"\n📄 Generated files:")
        print(f"   1. {raw_text_path.name} - Extracted text")
        print(f"   2. {json_path.name} - Structured data (for developers)")
        print(f"   3. {md_path.name} - Markdown report (for technical review)")
        print(f"   4. {pdf_path.name} - Beautiful PDF (FOR STAKEHOLDERS) ⭐")

        print(f"\n🎯 Share with stakeholders: {pdf_path}")
        print(f"\n💡 Tip: Open the PDF to see award-winning professional formatting!")

        print("\n✨ Advantages of this enhanced version:")
        print("   ✅ Beautiful PDF with excellent UI/UX")
        print("   ✅ Color-coded severity levels")
        print("   ✅ Professional formatting for non-technical readers")
        print("   ✅ Organized output in timestamped folders")
        print("   ✅ Multiple formats (PDF, JSON, Markdown, Raw Text)")
        print("   ✅ 100% offline - no data leaves your machine")

        print("\n")

    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        print("\n💡 Troubleshooting:")
        print("   1. Is LM Studio running?")
        print("   2. Is a model loaded in LM Studio?")
        print("   3. Is the local server started?")
        print(f"   4. Can you access: {lm_studio_url}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
