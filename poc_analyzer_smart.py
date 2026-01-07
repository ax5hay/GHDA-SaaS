#!/usr/bin/env python3
"""
SMART Government Health Data Analyzer
Automatically detects best available model from LM Studio
"""

import os
import sys
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

try:
    import docx
    from PyPDF2 import PdfReader
except ImportError:
    print("ERROR: Required packages not installed.")
    print("Install with: pip install openai python-docx PyPDF2 python-dateutil")
    sys.exit(1)

try:
    from openai import OpenAI
except ImportError:
    print("ERROR: OpenAI package not installed.")
    print("Install with: pip install openai")
    sys.exit(1)


def get_model_score(model_id: str) -> int:
    """Score models for suitability (higher = better)."""
    model_lower = model_id.lower()
    score = 0

    # Penalize reasoning models heavily
    if 'reasoning' in model_lower or 'o1' in model_lower:
        return -1000

    # Prefer specific models
    if 'qwen' in model_lower:
        score += 100
        if '14b' in model_lower or '32b' in model_lower:
            score += 50
    elif 'llama' in model_lower:
        score += 80
        if '8b' in model_lower or '70b' in model_lower:
            score += 30
    elif 'mistral' in model_lower:
        score += 70
        if 'instruct' in model_lower:
            score += 20

    # Require instruct/chat models
    if 'instruct' in model_lower or 'chat' in model_lower:
        score += 50
    else:
        score -= 100  # Penalize base models

    # Penalize very small models
    if '3b' in model_lower or '1b' in model_lower:
        score -= 30

    return score


def select_best_model(base_url: str) -> Optional[str]:
    """Automatically select best available model."""
    try:
        client = OpenAI(base_url=base_url, api_key="test")
        models_response = client.models.list()

        if not models_response.data:
            return None

        # Score and sort models
        scored_models = [
            (model.id, get_model_score(model.id))
            for model in models_response.data
        ]
        scored_models.sort(key=lambda x: x[1], reverse=True)

        # Return best model
        best_model = scored_models[0]
        return best_model[0] if best_model[1] > 0 else None

    except:
        return None


class SmartGovernmentHealthAnalyzer:
    """Smart analyzer with automatic model selection."""

    def __init__(self, base_url: str = "http://192.168.56.1:1234/v1", model_name: str = None):
        """Initialize with smart model selection."""
        self.base_url = base_url
        self.client = OpenAI(base_url=base_url, api_key="lm-studio")

        # Smart model selection
        if model_name:
            self.model_name = model_name
            print(f"🤖 Using specified model: {self.model_name}")
        else:
            env_model = os.environ.get("LM_STUDIO_MODEL")
            if env_model:
                self.model_name = env_model
                print(f"🤖 Using model from environment: {self.model_name}")
            else:
                print("🔍 Auto-detecting best available model...")
                best_model = select_best_model(base_url)
                if best_model:
                    self.model_name = best_model
                    print(f"✅ Auto-selected: {self.model_name}")
                else:
                    print("⚠️  Could not auto-detect model, using default")
                    self.model_name = "phi-4"

        # Warn about reasoning models
        if 'reasoning' in self.model_name.lower():
            print("\n" + "="*80)
            print("⚠️  WARNING: REASONING MODEL DETECTED")
            print("="*80)
            print(f"Model: {self.model_name}")
            print("\n❌ Reasoning models are NOT suitable for data extraction!")
            print("   • They are designed for complex logic problems")
            print("   • They take 5-30 minutes per document (or hang)")
            print("   • They often timeout before completing")
            print("\n💡 Recommended action:")
            print("   1. Cancel this run (Ctrl+C)")
            print("   2. Run: python select_model.py")
            print("   3. Choose a non-reasoning model (Qwen, Llama, Mistral)")
            print("\n" + "="*80)

            response = input("\nDo you want to continue anyway? (yes/no): ").strip().lower()
            if response != 'yes':
                print("Aborting...")
                sys.exit(0)

        print(f"✅ Connected to: {self.base_url}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        doc = docx.Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                text_parts.append(" | ".join(row_data))

        return "\n".join(text_parts)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)

        return "\n".join(text_parts)

    def extract_text(self, file_path: str) -> str:
        """Extract text from document."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension == ".docx":
            return self.extract_text_from_docx(str(file_path))
        elif extension == ".pdf":
            return self.extract_text_from_pdf(str(file_path))
        else:
            return file_path.read_text(encoding="utf-8", errors="ignore")

    def analyze_document(self, raw_text: str) -> Dict[str, Any]:
        """Analyze document with smart timeout and error handling."""

        # Simplified prompt for better results
        analysis_prompt = f"""Extract structured data from this health clinic report.

REPORT:
{raw_text[:15000]}

Return ONLY valid JSON with this structure:
{{
  "facility": {{"name": "...", "type": "...", "district": "..."}},
  "clinic_date": "YYYY-MM-DD",
  "beneficiaries": {{"expected": N, "attended": N}},
  "executive_summary": "2-3 sentence summary",
  "key_findings": ["finding 1", "finding 2"],
  "overall_score": 0-100
}}"""

        print(f"\n📏 Document length: {len(raw_text)} characters")

        if len(raw_text) > 15000:
            print(f"⚠️  Document truncated to 15,000 chars for processing")

        print("🤖 Analyzing with local AI...")

        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Extract data as JSON only."},
                    {"role": "user", "content": analysis_prompt}
                ],
                temperature=0.3,
                max_tokens=4000,
                timeout=180.0  # 3 minute timeout
            )

            response_text = response.choices[0].message.content

            if not response_text:
                raise ValueError("Empty response from model")

            # Extract JSON
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                response_text = json_match.group(0)

            response_text = re.sub(r',(\s*[}\]])', r'\1', response_text)

            return json.loads(response_text)

        except Exception as e:
            print(f"❌ Analysis failed: {e}")
            return {
                "error": str(e),
                "facility": {"name": "Unknown", "type": "Unknown"},
                "executive_summary": "Analysis failed",
                "overall_score": 0
            }

    def generate_report(self, analysis: Dict[str, Any], output_path: str):
        """Generate markdown report."""
        report = f"""# Government Health Data Analysis

**Generated**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
**Model**: {self.model_name}

## Summary

{analysis.get('executive_summary', 'No summary available')}

**Overall Score**: {analysis.get('overall_score', 'N/A')}/100

## Key Findings

"""
        for i, finding in enumerate(analysis.get('key_findings', []), 1):
            report += f"{i}. {finding}\n"

        report += f"\n\n---\n\n**Facility**: {analysis.get('facility', {}).get('name', 'Unknown')}\n"
        report += f"**Clinic Date**: {analysis.get('clinic_date', 'Unknown')}\n"

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report)

        json_path = output_path.replace('.md', '.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, indent=2, ensure_ascii=False)

        print(f"✅ Report: {output_path}")
        print(f"✅ JSON: {json_path}")


def main():
    """Main entry point."""

    print("""
╔══════════════════════════════════════════════════════════════════╗
║  🏥 SMART GOVERNMENT HEALTH DATA ANALYZER 🏥                    ║
║                                                                  ║
║  Automatic Model Selection + Smart Error Handling               ║
╚══════════════════════════════════════════════════════════════════╝
""")

    if len(sys.argv) != 2:
        print("Usage: python poc_analyzer_smart.py <path_to_report>")
        print("\n💡 Tips:")
        print("   • Model is auto-selected (or set LM_STUDIO_MODEL)")
        print("   • Run 'python select_model.py' to choose model manually")
        sys.exit(1)

    input_file = sys.argv[1]

    if not Path(input_file).exists():
        print(f"❌ File not found: {input_file}")
        sys.exit(1)

    try:
        analyzer = SmartGovernmentHealthAnalyzer()

        print(f"\n📄 Processing: {input_file}")
        raw_text = analyzer.extract_text(input_file)
        print(f"✅ Extracted {len(raw_text)} characters")

        analysis = analyzer.analyze_document(raw_text)

        output_path = Path(input_file).stem + "_SMART_ANALYSIS.md"
        analyzer.generate_report(analysis, output_path)

        print(f"\n🎉 Analysis complete!")

    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
